# -*- coding: utf-8 -*-
"""
Facebook 貼文 / 相片 / 短影片 全自動擷取（最終交付版）

功能：
- 支援 share / permalink / photo.php / reel
- 文字去 FB UI 雜訊
- 圖片直接嵌入 DOCX（不落地）
- 若有影片：下載 mp4（不轉畫質）
- 嘗試用 ffmpeg remux 修復「下載但無法播放」的影片
- DOCX / MP4 同名
- 依社團或來源名稱建資料夾
- 全程無 Enter 互動，完成即回 CMD
"""

import re
import time
import shutil
import subprocess
from io import BytesIO
from pathlib import Path
from datetime import datetime
from urllib.parse import urlsplit, urlunsplit

from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Pt, Inches, Cm
from PIL import Image


# ===================== 基本設定 =====================
BASE_DIR = Path(r"F:\F\AI")
PROFILE_DIR = BASE_DIR / "_fb_profile"


# ===================== 工具 =====================
def normalize_url(u: str) -> str:
    sp = urlsplit((u or "").strip())
    return urlunsplit((sp.scheme, sp.netloc, sp.path, sp.query, ""))


def safe_filename(s: str, max_len=80) -> str:
    s = re.sub(r'[<>:"/\\|?*]', "_", s or "")
    s = re.sub(r"\s+", " ", s).strip()
    return (s[:max_len] or "FB內容")


def choose_available_path(folder: Path, base: str, suffix: str) -> Path:
    p = folder / f"{base}{suffix}"
    if not p.exists():
        return p
    for i in range(1, 100):
        p2 = folder / f"{base}_{i:02d}{suffix}"
        if not p2.exists():
            return p2
    return folder / f"{base}_{int(time.time())}{suffix}"


def scroll(page, n=8):
    for _ in range(n):
        page.mouse.wheel(0, 1600)
        page.wait_for_timeout(600)


# ===================== FB UI 雜訊過濾 =====================
_UI = [
    r"^讚$|^留言$|^分享$|^回覆$|查看更多|See more|翻譯|See translation",
    r"^\d+人讚$|^\d+次分享$|^\d+則留言$",
]

def clean_text(t: str) -> str:
    if not t:
        return ""
    out = []
    for line in t.replace("\r", "").split("\n"):
        line = line.strip()
        if not line or len(line) < 3:
            continue
        if any(re.search(p, line, re.I) for p in _UI):
            continue
        if line not in out:
            out.append(line)
    return "\n".join(out)


def first_line(t: str) -> str:
    for l in t.splitlines():
        if l.strip():
            return l.strip()
    return ""


# ===================== 來源名稱 =====================
def get_source(page) -> str:
    for sel in ['a[href*="/groups/"] span', 'a[href*="/groups/"]', 'h1', 'title']:
        try:
            e = page.locator(sel)
            if e.count():
                return e.first.inner_text().strip()[:60]
        except:
            pass
    return "Facebook"


# ===================== 圖片 =====================
def collect_images(ctx, container):
    imgs = []
    seen = set()
    for img in container.locator("img").all():
        try:
            src = img.get_attribute("src")
            if not src or "emoji" in src or src in seen:
                continue
            seen.add(src)
            r = ctx.request.get(src)
            im = Image.open(BytesIO(r.body())).convert("RGB")
            bio = BytesIO()
            im.save(bio, "JPEG", quality=92)
            bio.seek(0)
            imgs.append(bio)
        except:
            pass
    return imgs


# ===================== 影片修復（remux） =====================
def try_remux_mp4(mp4: Path) -> bool:
    if not shutil.which("ffmpeg"):
        return False
    fixed = mp4.with_name(mp4.stem + "_fixed.mp4")
    cmd = [
        "ffmpeg", "-y",
        "-i", str(mp4),
        "-c", "copy",
        "-movflags", "+faststart",
        str(fixed)
    ]
    try:
        subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, check=True)
        if fixed.exists() and fixed.stat().st_size > mp4.stat().st_size * 0.5:
            mp4.unlink()
            fixed.rename(mp4)
            return True
    except:
        pass
    return False


# ===================== 主流程 =====================
def main():
    url = normalize_url(input("請輸入 FB 網址：\n"))

    BASE_DIR.mkdir(parents=True, exist_ok=True)
    PROFILE_DIR.mkdir(parents=True, exist_ok=True)

    with sync_playwright() as p:
        ctx = p.chromium.launch_persistent_context(
            user_data_dir=str(PROFILE_DIR),
            headless=False,
            locale="zh-TW",
            viewport={"width": 1280, "height": 900},
        )

        page = ctx.new_page()

        video_candidates = []

        def on_resp(resp):
            try:
                if ".mp4" in resp.url.lower():
                    video_candidates.append(resp.url)
            except:
                pass

        page.on("response", on_resp)

        page.goto(url, timeout=60000)
        scroll(page, 10)
        page.wait_for_timeout(2000)

        source = get_source(page)
        folder = BASE_DIR / safe_filename(source)
        folder.mkdir(exist_ok=True)

        # ===== 抓內容 =====
        container = page.locator('div[role="article"]').first
        if not container.count():
            container = page.locator('div[role="main"]').first

        text = clean_text(container.inner_text())
        title = page.title() or "Facebook 內容"

        # ===== DOCX =====
        today = datetime.now().strftime("%Y%m%d")
        base = safe_filename(first_line(text) or title)
        base_name = f"{today}_{base}"

        docx_path = choose_available_path(folder, base_name, ".docx")

        doc = Document()
        sec = doc.sections[0]
        for m in [sec.top_margin, sec.bottom_margin, sec.left_margin, sec.right_margin]:
            m = Cm(1.5)

        doc.add_paragraph(f"【來源】{source}").runs[0].bold = True
        doc.add_heading(title, 1)

        for l in text.splitlines():
            doc.add_paragraph(l)

        imgs = collect_images(ctx, container)
        for bio in imgs:
            doc.add_picture(bio, width=Inches(5.8))

        doc.save(docx_path)
        print("📄 DOCX 完成：", docx_path)

        # ===== 影片 =====
        if video_candidates:
            mp4_path = docx_path.with_suffix(".mp4")
            try:
                r = ctx.request.get(video_candidates[0], timeout=600000)
                mp4_path.write_bytes(r.body())
                print("🎬 影片下載完成：", mp4_path)

                if try_remux_mp4(mp4_path):
                    print("🎬 已重新封裝（可播放）")
                else:
                    print("🎬 未修復（可能仍可播放，或無 ffmpeg）")
            except:
                print("⚠️ 影片下載失敗")
        else:
            print("🎬 本貼文無影片")

        ctx.close()


if __name__ == "__main__":
    main()
