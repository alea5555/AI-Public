# 檔名：P爬文章批次轉成docx_痞客邦.py
import os
import re
import time
import json
import csv
import traceback
from io import BytesIO
from urllib.parse import urljoin, urlparse
from datetime import datetime

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from docx.image.exceptions import UnrecognizedImageError

# 可選：用來把 webp 轉 png（沒裝也沒關係，會自動跳過）
try:
    from PIL import Image
    PIL_OK = True
except Exception:
    PIL_OK = False


# ===== 你環境的路徑 =====
CSV_PATH = r"F:\F\AI\web\web.csv"   # A欄=網址，B欄=名稱(可空)
OUT_DIR  = r"F:\F\AI\web"          # docx 輸出資料夾
SLEEP_SEC = 0.5                    # 下載圖片間隔


# ========== 單次版：safe_filename（同邏輯） ==========
def safe_filename(name: str, max_len: int = 120) -> str:
    name = re.sub(r'[<>:"/\\|?*]', "_", (name or "")).strip()
    name = re.sub(r"\s+", " ", name)
    if len(name) > max_len:
        name = name[:max_len].rstrip()
    return name or "output"


# ========== 讀 CSV：支援 Excel 常見編碼 ==========
def read_urls_from_csv(csv_path: str):
    encodings_to_try = ["utf-8-sig", "cp950", "big5", "utf-8", "latin-1"]
    last_err = None

    for enc in encodings_to_try:
        try:
            items = []
            with open(csv_path, "r", encoding=enc, newline="") as f:
                reader = csv.reader(f)
                for r in reader:
                    if not r:
                        continue
                    url = (r[0] if len(r) > 0 else "").strip()
                    name = (r[1] if len(r) > 1 else "").strip()

                    # 跳過表頭
                    if url.lower() in {"url", "網址"}:
                        continue

                    # 空白行跳過（不要 break）
                    if not url:
                        continue

                    items.append((url, name))
            return items
        except UnicodeDecodeError as e:
            last_err = e
            continue

    raise last_err


# ========== 單次版：fetch_html（同邏輯） ==========
def fetch_html(session: requests.Session, url: str) -> str:
    url = url.split("#", 1)[0]
    r = session.get(url, timeout=30)
    r.raise_for_status()
    if not r.encoding or r.encoding.lower() == "iso-8859-1":
        r.encoding = r.apparent_encoding or "utf-8"
    return r.text


# =========================
# ✅ 單次版：選「最像正文」的容器
# =========================
def _node_score(node) -> int:
    if not node:
        return -10**9

    txt = node.get_text(" ", strip=True)
    tlen = len(txt)

    p = len(node.find_all("p"))
    li = len(node.find_all("li"))
    h = len(node.find_all(["h1", "h2", "h3", "h4"]))
    pre = len(node.find_all("pre"))
    code = len(node.find_all("code"))
    bq = len(node.find_all("blockquote"))
    img = len(node.find_all("img"))

    bad = 0
    for bad_sel in ["nav", "header", "footer", "aside"]:
        bad += len(node.find_all(bad_sel))

    cls = " ".join(node.get("class", [])).lower()
    nid = (node.get("id") or "").lower()
    if any(k in cls for k in ["comment", "sidebar", "related", "recommend", "widget", "breadcrumb", "footer"]):
        bad += 10
    if any(k in nid for k in ["comment", "sidebar", "related", "recommend", "footer"]):
        bad += 10

    score = 0
    score += min(tlen, 20000)
    score += p * 300
    score += li * 120
    score += h * 200
    score += pre * 200
    score += code * 50
    score += bq * 150
    score += img * 10
    score -= bad * 500
    return score


def pick_content_root(soup: BeautifulSoup):
    selectors = [
        "article",
        "main",
        ".vditor-reset",
        ".markdown-body",
        ".post-detail",
        ".post-content",
        ".entry-content",
        ".article-content",
        ".content",
        "#content",
        "#__next",
        "body",
    ]

    candidates = []
    for sel in selectors:
        for node in soup.select(sel):
            candidates.append(node)

    if not candidates:
        return soup.body or soup

    best = max(candidates, key=_node_score)
    return best


def is_probably_nav_or_junk(tag) -> bool:
    if tag.name in {"nav", "header", "footer", "aside", "script", "style", "noscript"}:
        return True
    cls = " ".join(tag.get("class", [])).lower()
    if any(k in cls for k in ["share", "related", "sidebar", "widget", "comment", "ads", "advert", "breadcrumb"]):
        return True
    return False


def iter_content_blocks(root):
    for t in root.find_all(["script", "style", "noscript"]):
        t.decompose()

    for el in root.find_all(["h1", "h2", "h3", "h4", "p", "li", "blockquote", "pre", "img"]):
        if is_probably_nav_or_junk(el):
            continue

        if el.name in ["h1", "h2", "h3", "h4"]:
            txt = el.get_text(" ", strip=True)
            if txt:
                yield ("heading", el.name, txt)
            continue

        if el.name == "p":
            txt = el.get_text(" ", strip=True)
            if txt:
                yield ("p", txt)
            continue

        if el.name == "li":
            txt = el.get_text(" ", strip=True)
            if txt:
                yield ("li", txt)
            continue

        if el.name == "blockquote":
            txt = el.get_text(" ", strip=True)
            if txt:
                yield ("quote", txt)
            continue

        if el.name == "pre":
            txt = el.get_text("\n", strip=True)
            if txt:
                yield ("codeblock", txt)
            continue

        if el.name == "img":
            src = el.get("src") or el.get("data-src") or el.get("data-lazy-src") or el.get("data-original")
            if not src:
                continue
            alt = (el.get("alt") or "").strip()
            yield ("img", src, alt)
            continue


def download_image(session: requests.Session, img_url: str):
    try:
        r = session.get(img_url, timeout=30)
        r.raise_for_status()
        ctype = (r.headers.get("Content-Type") or "").lower()
        if "image" not in ctype:
            return None, ctype
        return r.content, ctype
    except Exception:
        return None, ""


# --------- 單次版保底：從 Next/Nuxt JSON 找正文 ---------
def _extract_next_data_json(html: str):
    m = re.search(r'<script[^>]+id="__NEXT_DATA__"[^>]*>(.*?)</script>', html, re.S | re.I)
    if not m:
        return None
    try:
        return json.loads(m.group(1).strip())
    except Exception:
        return None


def _extract_nuxt_data_json(html: str):
    m = re.search(r'window\.__NUXT__\s*=\s*(\{.*?\});\s*</script>', html, re.S | re.I)
    if not m:
        m = re.search(r'window\.__NUXT__\s*=\s*(\{.*\})\s*;?', html, re.S | re.I)
    if not m:
        return None
    raw = m.group(1).strip()
    raw = re.sub(r"\bundefined\b", "null", raw)
    try:
        return json.loads(raw)
    except Exception:
        return None


def _score_candidate(key: str, s: str) -> int:
    k = (key or "").lower()
    score = 0
    if any(x in k for x in ["content", "html", "markdown", "body", "article", "text"]):
        score += 200
    if "<p" in s or "<h" in s or "</" in s:
        score += 150
    if "\n" in s:
        score += 50
    if "function(" in s or "var " in s or "webpack" in s:
        score -= 300
    score += min(len(s) // 50, 400)
    return score


def _find_best_long_text(obj, key_path=""):
    best = ("", "", -10**9)

    def walk(x, kp):
        nonlocal best
        if isinstance(x, dict):
            for k, v in x.items():
                walk(v, f"{kp}.{k}" if kp else str(k))
        elif isinstance(x, list):
            for i, v in enumerate(x):
                walk(v, f"{kp}[{i}]")
        elif isinstance(x, str):
            s = x.strip()
            if len(s) < 200:
                return
            sc = _score_candidate(kp.split(".")[-1], s)
            if sc > best[2]:
                best = (kp, s, sc)

    walk(obj, key_path)
    return best


def try_extract_article_text_from_scripts(html: str):
    data = _extract_next_data_json(html)
    if not data:
        data = _extract_nuxt_data_json(html)
    if not data:
        return None

    kp, text, score = _find_best_long_text(data)
    if score < 0:
        return None
    return text


def add_plaintext_to_doc(doc: Document, text: str):
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    blocks = [b.strip() for b in re.split(r"\n{2,}", text) if b.strip()]
    for b in blocks:
        if re.match(r"^#{1,4}\s+", b):
            level = len(re.match(r"^(#+)", b).group(1))
            title = re.sub(r"^#{1,4}\s+", "", b).strip()
            doc.add_heading(title, level=min(level, 4))
        else:
            lines = b.split("\n")
            p = doc.add_paragraph(lines[0])
            for line in lines[1:]:
                p.add_run("\n" + line)


# =========================
# ✅ 單次版：抓建檔日期 → yyyymmdd
# =========================
def _parse_date_to_yyyymmdd(s: str):
    if not s:
        return None
    s = s.strip()

    m = re.search(r"(20\d{2})-(\d{1,2})-(\d{1,2})", s)
    if m:
        y, mo, d = m.group(1), int(m.group(2)), int(m.group(3))
        return f"{y}{mo:02d}{d:02d}"

    m = re.search(r"(20\d{2})/(\d{1,2})/(\d{1,2})", s)
    if m:
        y, mo, d = m.group(1), int(m.group(2)), int(m.group(3))
        return f"{y}{mo:02d}{d:02d}"

    return None


def extract_date8(soup: BeautifulSoup, html: str) -> str:
    meta_keys = [
        ("property", "article:published_time"),
        ("property", "og:published_time"),
        ("name", "pubdate"),
        ("name", "publishdate"),
        ("name", "publish_date"),
        ("name", "date"),
        ("itemprop", "datePublished"),
    ]

    for attr, val in meta_keys:
        tag = soup.find("meta", attrs={attr: val})
        if tag and tag.get("content"):
            d8 = _parse_date_to_yyyymmdd(tag["content"])
            if d8:
                return d8

    t = soup.find("time")
    if t:
        dt = t.get("datetime") or t.get_text(" ", strip=True)
        d8 = _parse_date_to_yyyymmdd(dt)
        if d8:
            return d8

    m = re.search(r"\b(20\d{2})[-/](\d{1,2})[-/](\d{1,2})\b", html)
    if m:
        y, mo, d = m.group(1), int(m.group(2)), int(m.group(3))
        return f"{y}{mo:02d}{d:02d}"

    return datetime.now().strftime("%Y%m%d")


# =========================
# ✅ 單次版：webp 轉 png（可選）
# =========================
def maybe_convert_webp_to_png_bytes(img_bytes: bytes, ctype: str, img_url: str):
    low_ct = (ctype or "").lower()
    ext = os.path.splitext(urlparse(img_url).path.lower())[1]

    is_webp = ("image/webp" in low_ct) or (ext == ".webp")
    if not is_webp:
        return None
    if not PIL_OK:
        return None

    try:
        im = Image.open(BytesIO(img_bytes))
        out = BytesIO()
        im.convert("RGB").save(out, format="PNG")
        return out.getvalue()
    except Exception:
        return None


# ========== 額外：標題清理（像你 JS：切掉 @ / :: 後綴） ==========
def clean_title_like_js(title: str) -> str:
    t = (title or "").strip()
    if not t:
        return ""
    t = re.split(r"\s*@\s*|\s*::\s*", t, maxsplit=1)[0].strip()
    return t


# =========================
# ✅ 批次：把「單次流程」包成一個函式
# =========================
def build_docx_for_one_url(session: requests.Session, url: str, name_from_csv: str):
    html = fetch_html(session, url)
    soup = BeautifulSoup(html, "lxml")

    # 檔名：B欄優先；B欄空白 → 用頁面 title
    if name_from_csv and name_from_csv.strip():
        file_base = safe_filename(name_from_csv.strip())
        page_title = name_from_csv.strip()
    else:
        page_title = soup.title.get_text(strip=True) if soup.title else "article"
        page_title = clean_title_like_js(page_title)
        file_base = safe_filename(page_title)

    # 日期：單次版 extract_date8
    date8 = extract_date8(soup, html)

    # 輸出檔名：YYYYMMDD_名稱.docx（你要的格式）
    out_path = os.path.join(OUT_DIR, f"{date8}_{file_base}.docx")

    # 內容：完全照單次版
    doc = Document()
    doc.add_heading(page_title, level=0)
    doc.add_paragraph(f"來源網址：{url}")
    doc.add_paragraph(f"建檔日期：{date8}")
    doc.add_paragraph("")

    root = pick_content_root(soup)

    img_count = 0
    text_count = 0

    blocks = list(iter_content_blocks(root))
    for block in blocks:
        kind = block[0]

        if kind == "heading":
            _, tagname, txt = block
            level_map = {"h1": 1, "h2": 2, "h3": 3, "h4": 4}
            doc.add_heading(txt, level=level_map.get(tagname, 2))
            text_count += 1

        elif kind == "p":
            _, txt = block
            doc.add_paragraph(txt)
            text_count += 1

        elif kind == "li":
            _, txt = block
            doc.add_paragraph(txt, style="List Bullet")
            text_count += 1

        elif kind == "quote":
            _, txt = block
            doc.add_paragraph(txt, style="Intense Quote")
            text_count += 1

        elif kind == "codeblock":
            _, txt = block
            p = doc.add_paragraph()
            run = p.add_run(txt)
            run.font.name = "Consolas"
            text_count += 1

        elif kind == "img":
            _, src, alt = block
            img_url = urljoin(url.split("#", 1)[0], src)

            path = urlparse(img_url).path.lower()
            if any(path.endswith(x) for x in [".svg", ".ico"]):
                continue

            img, ctype = download_image(session, img_url)
            if not img:
                continue

            if alt:
                doc.add_paragraph(alt)

            try:
                converted = maybe_convert_webp_to_png_bytes(img, ctype, img_url)
                if converted:
                    doc.add_picture(BytesIO(converted), width=Inches(6.0))
                else:
                    doc.add_picture(BytesIO(img), width=Inches(6.0))

                img_count += 1
                time.sleep(SLEEP_SEC)

            except UnrecognizedImageError:
                # 單次版也是跳過
                continue
            except Exception:
                continue

    # ✅ 保底：如果 DOM 幾乎抓不到文字，就從 script JSON 抽正文（單次版保底）
    if text_count <= 2:
        extracted = try_extract_article_text_from_scripts(html)
        if extracted:
            doc.add_page_break()
            doc.add_heading("（保底抽取內容）", level=1)

            if "<p" in extracted or "<h" in extracted or "</" in extracted:
                soup2 = BeautifulSoup(extracted, "lxml")
                root2 = pick_content_root(soup2)
                for block in iter_content_blocks(root2):
                    kind = block[0]
                    if kind == "heading":
                        _, tagname, txt = block
                        level_map = {"h1": 1, "h2": 2, "h3": 3, "h4": 4}
                        doc.add_heading(txt, level=level_map.get(tagname, 2))
                        text_count += 1
                    elif kind == "p":
                        _, txt = block
                        doc.add_paragraph(txt)
                        text_count += 1
                    elif kind == "li":
                        _, txt = block
                        doc.add_paragraph(txt, style="List Bullet")
                        text_count += 1
                    elif kind == "quote":
                        _, txt = block
                        doc.add_paragraph(txt, style="Intense Quote")
                        text_count += 1
                    elif kind == "codeblock":
                        _, txt = block
                        p = doc.add_paragraph()
                        run = p.add_run(txt)
                        run.font.name = "Consolas"
                        text_count += 1
            else:
                add_plaintext_to_doc(doc, extracted)
                text_count += 1

    doc.save(out_path)
    return out_path, text_count, img_count, date8, page_title


def main():
    print("[START] 單次版 → 批次版（完全沿用單次正文抽取/保底抽文/日期）")
    print(f"[INFO] CSV_PATH: {CSV_PATH}")
    print(f"[INFO] OUT_DIR : {OUT_DIR}")

    if not os.path.isfile(CSV_PATH):
        print(f"[ERROR] 找不到 CSV：{CSV_PATH}")
        return

    os.makedirs(OUT_DIR, exist_ok=True)

    items = read_urls_from_csv(CSV_PATH)
    print(f"[INFO] CSV 讀到 {len(items)} 筆")

    if not items:
        print("[WARN] CSV 沒有任何網址")
        return

    ok = 0
    skip = 0
    fail = 0

    with requests.Session() as s:
        # headers：沿用單次版那套（你單次能抓到內容就別亂改）
        s.headers.update({
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8",
            "Accept-Language": "zh-TW,zh;q=0.9,en;q=0.6",
            "Referer": "https://www.codefather.cn/",
        })

        for idx, (url, name_from_csv) in enumerate(items, start=1):
            try:
                # 先用 B 欄/頁面 title 算出輸出檔名，若存在就跳過
                # （為了保留你要的：同名就覆蓋 or 跳過？這裡採「存在就跳過」）
                # 若你要「覆蓋」我也可以改成直接寫入覆蓋。
                tmp_name = name_from_csv.strip() if name_from_csv else ""
                if tmp_name:
                    file_base = safe_filename(tmp_name)
                else:
                    # 先抓一次 title 來算檔名（避免每次都重抓）
                    html_peek = fetch_html(s, url)
                    soup_peek = BeautifulSoup(html_peek, "lxml")
                    page_title_peek = soup_peek.title.get_text(strip=True) if soup_peek.title else "article"
                    page_title_peek = clean_title_like_js(page_title_peek)
                    file_base = safe_filename(page_title_peek)
                    date8_peek = extract_date8(soup_peek, html_peek)
                    out_peek = os.path.join(OUT_DIR, f"{date8_peek}_{file_base}.docx")
                    if os.path.exists(out_peek):
                        print(f"[SKIP] ({idx}/{len(items)}) 已存在：{os.path.basename(out_peek)}")
                        skip += 1
                        continue
                    # 沒存在就直接用 peek 的 html 也行，但為了簡潔就下面正常跑一次 build

                # 正式跑（完全走單次流程）
                print(f"[DO] ({idx}/{len(items)}) {url}")
                out_path, text_count, img_count, date8, page_title = build_docx_for_one_url(s, url, name_from_csv)

                print(f"[OK]  {os.path.basename(out_path)} | 日期={date8} | 文字≈{text_count} | 圖片={img_count}")
                ok += 1

            except Exception as e:
                print(f"[ERR] ({idx}/{len(items)}) {url}\n      {e}")
                fail += 1
                continue

    print(f"\n[DONE] OK={ok}, SKIP={skip}, FAIL={fail}")


if __name__ == "__main__":
    try:
        main()
    except Exception:
        print("[FATAL] 未捕捉例外：")
        traceback.print_exc()
