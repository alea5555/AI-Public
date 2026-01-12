import os
import re
import time
import traceback
from io import BytesIO
from urllib.parse import urljoin, urlparse
from datetime import datetime

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from docx.image.exceptions import UnrecognizedImageError

# 可選：webp -> png（沒裝 PIL 也能跑，只是 webp 可能跳過）
try:
    from PIL import Image
    PIL_OK = True
except Exception:
    PIL_OK = False

OUT_DIR = r"F:\F\AI"
SLEEP_SEC = 0.35
MAX_MEDIA = 40


def safe_filename(name: str, max_len: int = 120) -> str:
    name = re.sub(r'[<>:"/\\|?*]', "_", str(name)).strip()
    name = re.sub(r"\s+", "_", name)          # ✅ 空白 -> _
    name = re.sub(r"_+", "_", name)           # 合併多個 _
    if len(name) > max_len:
        name = name[:max_len].rstrip("_")
    return name or "output"


def _browser_headers(referer: str = "https://www.threads.com/") -> dict:
    return {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/122.0.0.0 Safari/537.36"
        ),
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8",
        "Accept-Language": "zh-TW,zh;q=0.9,en;q=0.8",
        "Referer": referer,
        "Connection": "keep-alive",
        "Upgrade-Insecure-Requests": "1",
    }


def parse_threads_url_info(url: str) -> str:
    u = urlparse(url)
    parts = [p for p in u.path.split("/") if p]
    handle = ""
    post_id = ""
    # 常見：/@handle/post/POSTID
    for i, p in enumerate(parts):
        if p.startswith("@"):
            handle = p.lstrip("@")
        if p == "post" and i + 1 < len(parts):
            post_id = parts[i + 1]
    if handle and post_id:
        return f"threads_{handle}_{post_id}"
    if post_id:
        return f"threads_{post_id}"
    return "threads_post"


def fetch_html_requests(url: str, timeout: int = 30) -> str:
    url = url.split("#", 1)[0]
    with requests.Session() as s:
        s.headers.update(_browser_headers())
        r = s.get(url, timeout=timeout, allow_redirects=True)
        if not r.encoding or r.encoding.lower() == "iso-8859-1":
            r.encoding = r.apparent_encoding or "utf-8"
        return r.text


def fetch_playwright_bundle(url: str):
    """
    ✅ 用 Playwright 抓：
    - 可視 DOM 正文（用「最長可視文字區塊」策略）
    - 貼文區塊內圖片（currentSrc/src/srcset）
    - 整頁截圖（保底）
    """
    from playwright.sync_api import sync_playwright

    headless = os.environ.get("HEADLESS", "1") != "0"

    with sync_playwright() as p:
        browser = p.chromium.launch(
            headless=headless,
            args=[
                "--no-sandbox",
                "--disable-setuid-sandbox",
                "--disable-dev-shm-usage",
                "--disable-blink-features=AutomationControlled",
            ],
        )
        context = browser.new_context(
            user_agent=_browser_headers()["User-Agent"],
            locale="zh-TW",
            viewport={"width": 1366, "height": 900},
        )
        page = context.new_page()

        page.goto(url, wait_until="domcontentloaded", timeout=60000)
        page.wait_for_timeout(2200)

        # 觸發 lazy-load
        try:
            page.mouse.wheel(0, 1800)
            page.wait_for_timeout(1200)
        except Exception:
            pass

        # ✅ 1) 先抓「可視正文」：找 main 內最長的可視文字區塊
        post_text = ""
        try:
            post_text = page.evaluate(
                """() => {
                    const norm = (s) => (s || '')
                        .replace(/\\r/g, '')
                        .replace(/\\n{3,}/g, '\\n\\n')
                        .trim();

                    const isBad = (el) => {
                        if (!el) return true;
                        const tag = (el.tagName || '').toLowerCase();
                        if (['script','style','noscript','svg'].includes(tag)) return true;
                        const role = (el.getAttribute && el.getAttribute('role')) || '';
                        if (role && ['navigation','banner','dialog'].includes(role)) return true;
                        return false;
                    };

                    const main = document.querySelector('main') || document.body;
                    const cand = [];
                    const walker = document.createTreeWalker(main, NodeFilter.SHOW_ELEMENT, null);
                    while (walker.nextNode()) {
                        const el = walker.currentNode;
                        if (isBad(el)) continue;

                        // 排除明顯 UI：按鈕、輸入等
                        const tag = (el.tagName || '').toLowerCase();
                        if (['button','input','textarea','select'].includes(tag)) continue;

                        // 只取可見元素
                        const st = window.getComputedStyle(el);
                        if (!st || st.display === 'none' || st.visibility === 'hidden') continue;

                        const txt = norm(el.innerText);
                        if (!txt) continue;

                        // 文字太短不要
                        if (txt.length < 80) continue;

                        // 避免抓到整頁：只保留包含換行/段落的
                        const score = txt.length + (txt.includes('\\n') ? 120 : 0);

                        cand.push({score, txt});
                    }

                    cand.sort((a,b) => b.score - a.score);

                    // 取第一個最像正文的（通常就是貼文）
                    if (cand.length) {
                        // 再做一點過濾：不要包含太多網站固定字
                        const blacklist = ['Meta.ai', 'Cookie', '登入', 'Log in', 'Sign up'];
                        for (const c of cand) {
                            const bad = blacklist.some(k => c.txt.includes(k));
                            if (!bad) return c.txt;
                        }
                        return cand[0].txt;
                    }

                    return '';
                }"""
            )
        except Exception:
            post_text = ""

        # ✅ 2) 再抓「貼文區塊內圖片」：先鎖定 main，再抓所有 images 的 currentSrc/srcset
        img_urls = []
        try:
            img_urls = page.evaluate(
                """() => {
                    const out = new Set();
                    const main = document.querySelector('main') || document.body;

                    const imgs = Array.from(main.querySelectorAll('img'));
                    for (const img of imgs) {
                        const w = img.naturalWidth || 0;
                        const h = img.naturalHeight || 0;

                        // 過濾：太小的多半是頭像/圖示
                        if (w && h && (w < 120 || h < 120)) continue;

                        if (img.currentSrc) out.add(img.currentSrc);
                        if (img.src) out.add(img.src);

                        const ss = img.getAttribute('srcset');
                        if (ss) {
                            ss.split(',')
                              .map(s => s.trim().split(' ')[0])
                              .forEach(u => { if (u) out.add(u); });
                        }
                    }

                    // meta 圖
                    const og = document.querySelector('meta[property="og:image"]');
                    if (og && og.content) out.add(og.content);

                    return Array.from(out);
                }"""
            )
        except Exception:
            img_urls = []

        # ✅ 3) 整頁截圖保底
        screenshot_bytes = None
        try:
            screenshot_bytes = page.screenshot(full_page=True)
        except Exception:
            screenshot_bytes = None

        html = page.content()
        browser.close()

    # 過濾：data: / svg / ico
    cleaned = []
    seen = set()
    for u in (img_urls or []):
        u = (u or "").strip()
        if not u or u.startswith("data:"):
            continue
        pth = urlparse(u).path.lower()
        if pth.endswith(".svg") or pth.endswith(".ico"):
            continue
        if u in seen:
            continue
        seen.add(u)
        cleaned.append(u)

    return html, cleaned, screenshot_bytes, (post_text or "").strip()


def extract_meta_title_and_date(html: str):
    soup = BeautifulSoup(html, "lxml")

    def meta(prop=None, name=None):
        if prop:
            t = soup.find("meta", attrs={"property": prop})
            if t and t.get("content"):
                return t["content"].strip()
        if name:
            t = soup.find("meta", attrs={"name": name})
            if t and t.get("content"):
                return t["content"].strip()
        return ""

    title = meta(prop="og:title") or meta(name="twitter:title") or ""

    # 日期：拿不到就今天
    date8 = ""
    for prop in ["article:published_time", "og:published_time", "og:updated_time"]:
        t = meta(prop=prop)
        if t:
            m = re.search(r"(20\d{2})-(\d{1,2})-(\d{1,2})", t)
            if m:
                y, mo, d = m.group(1), int(m.group(2)), int(m.group(3))
                date8 = f"{y}{mo:02d}{d:02d}"
                break
    if not date8:
        date8 = datetime.now().strftime("%Y%m%d")

    return title.strip(), date8


def download_image_bytes(session: requests.Session, img_url: str):
    try:
        r = session.get(img_url, timeout=30)
        r.raise_for_status()
        ctype = (r.headers.get("Content-Type") or "").lower()
        if "image" not in ctype:
            return None, ctype
        return r.content, ctype
    except Exception:
        return None, ""


def maybe_convert_webp_to_png_bytes(img_bytes: bytes, ctype: str, img_url: str):
    low_ct = (ctype or "").lower()
    ext = os.path.splitext(urlparse(img_url).path.lower())[1]
    is_webp = ("image/webp" in low_ct) or (ext == ".webp")
    if not is_webp or not PIL_OK:
        return None
    try:
        im = Image.open(BytesIO(img_bytes))
        out = BytesIO()
        im.convert("RGB").save(out, format="PNG")
        return out.getvalue()
    except Exception:
        return None


def add_picture_to_doc(doc: Document, img_bytes: bytes, width_inches: float = 6.3) -> bool:
    try:
        doc.add_picture(BytesIO(img_bytes), width=Inches(width_inches))
        return True
    except UnrecognizedImageError:
        return False
    except Exception:
        return False


def main():
    url = input("請輸入 Threads 貼文網址：\n").strip()
    if not url:
        print("❌ 未輸入網址，結束")
        return

    os.makedirs(OUT_DIR, exist_ok=True)

    # ✅ Threads：直接用 Playwright bundle（最穩）
    html, img_urls, screenshot_bytes, dom_text = fetch_playwright_bundle(url)

    meta_title, date8 = extract_meta_title_and_date(html)
    fallback_title = parse_threads_url_info(url)

    # ✅ 標題：meta_title 有用就用，否則用網址推導
    title = meta_title if meta_title and meta_title.lower() != "threads" else fallback_title

    # ✅ 正文：優先用 DOM 抽到的可視文字（你現在缺的就是這段）
    text = dom_text.strip()

    out_path = os.path.join(OUT_DIR, f"{date8}_{safe_filename(title)}.docx")

    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"來源網址：{url}")
    doc.add_paragraph(f"建檔日期：{date8}")
    doc.add_paragraph("")

    if text:
        doc.add_paragraph(text)
    else:
        doc.add_paragraph("（未成功抽取到正文，可能貼文權限受限或需登入）")

    # ✅ 插入圖片（從 DOM 抓到的 img）
    img_urls = (img_urls or [])[:MAX_MEDIA]
    img_count = 0

    with requests.Session() as s:
        s.headers.update(_browser_headers(referer="https://www.threads.com/"))

        for img_url in img_urls:
            img, ctype = download_image_bytes(s, img_url)
            if not img:
                continue

            converted = maybe_convert_webp_to_png_bytes(img, ctype, img_url)
            ok = False
            if converted:
                ok = add_picture_to_doc(doc, converted)
            if not ok:
                ok = add_picture_to_doc(doc, img)

            if ok:
                img_count += 1
                time.sleep(SLEEP_SEC)

    # ✅ 如果仍然沒抓到任何圖片：插入整頁截圖（保底一定有）
    if img_count == 0 and screenshot_bytes:
        doc.add_page_break()
        doc.add_heading("貼文截圖", level=1)
        if add_picture_to_doc(doc, screenshot_bytes, width_inches=6.8):
            img_count = 1

    doc.save(out_path)
    print(f"✅ 完成：{out_path}")
    print(f"📌 圖片：{img_count} 張")


if __name__ == "__main__":
    try:
        main()
    except Exception:
        print("❌ 程式發生未處理例外：")
        traceback.print_exc()
        # ✅ 不暫停，直接回 CMD
