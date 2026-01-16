import asyncio
import os
import re
import time
from datetime import datetime
from io import BytesIO
from urllib.parse import urljoin, urlparse

# --- docx ---
from docx import Document
from docx.shared import Inches

# --- http ---
import requests

# --- playwright ---
from playwright.async_api import async_playwright


def safe_filename(name: str, max_len: int = 120) -> str:
    name = (name or "").strip()
    if not name:
        return "output"
    name = re.sub(r'[<>:"/\\|?*\x00-\x1F]', "_", name)
    name = re.sub(r"\s+", " ", name).strip()
    if len(name) > max_len:
        name = name[:max_len].strip()
    return name or "output"


def sanitize_text(s: str) -> str:
    if s is None:
        return ""
    # 清掉 XML 1.0 不允許控制字元，避免 Word 進修復流程
    out_chars = []
    for ch in str(s):
        code = ord(ch)
        if code in (0x09, 0x0A, 0x0D):  # tab, lf, cr
            out_chars.append(ch)
            continue
        if code < 0x20:
            continue
        if 0xD800 <= code <= 0xDFFF:
            continue
        if code in (0xFFFE, 0xFFFF):
            continue
        out_chars.append(ch)

    out = "".join(out_chars)
    out = out.replace("\u00a0", " ")
    out = re.sub(r"\s+", " ", out).strip()
    return out


def uniq_keep_order(items):
    seen = set()
    out = []
    for x in items:
        x = (x or "").strip()
        if not x:
            continue
        if x in seen:
            continue
        seen.add(x)
        out.append(x)
    return out


def normalize_paragraphs(lines):
    # 去掉常見雜訊（你可自行加）
    junk_contains = ["MULTISTRATEGYS", "策略市集", "策略生成器"]
    cleaned = []
    for s in lines:
        s = sanitize_text(s)
        if not s or len(s) < 2:
            continue
        if any(j in s for j in junk_contains):
            continue
        cleaned.append(s)

    cleaned = uniq_keep_order(cleaned)

    # 合併短碎片
    merged = []
    for s in cleaned:
        if not merged:
            merged.append(s)
            continue
        prev = merged[-1]
        if len(prev) < 12 and len(s) < 20 and not re.search(r"[。！？:：]$", prev):
            merged[-1] = sanitize_text(prev + s)
        else:
            merged.append(s)

    # 切段
    paras = []
    buf = []
    def flush():
        nonlocal buf
        t = sanitize_text(" ".join(buf))
        if t:
            paras.append(t)
        buf = []

    for s in merged:
        is_heading_like = bool(re.match(r"^([一二三四五六七八九十]、|[0-9]+[.)])", s)) or (len(s) <= 28 and "：" in s)
        if is_heading_like and buf:
            flush()

        buf.append(s)

        cur = " ".join(buf)
        if re.search(r"[。！？]$", s) or len(cur) > 260:
            flush()

    flush()
    return paras


def safe_source_url(url: str) -> str:
    # 避免 Word 自動當外部連結觸發更新提示
    return url.replace("https://", "hxxps://").replace("http://", "hxxp://")


def fetch_image_bytes(url: str, timeout: int = 30):
    try:
        r = requests.get(url, timeout=timeout, headers={"User-Agent": "Mozilla/5.0"})
        if r.status_code != 200:
            return None, ""
        ct = (r.headers.get("content-type") or "").lower()
        return r.content, ct
    except Exception:
        return None, ""


def try_convert_to_png(image_bytes: bytes):
    """
    盡量把圖片轉 PNG（最穩）
    - 若 PIL 不可用或轉換失敗，就回傳 None
    """
    try:
        from PIL import Image  # noqa
    except Exception:
        return None

    try:
        im = Image.open(BytesIO(image_bytes))
        # 取第一幀（gif）
        if getattr(im, "is_animated", False):
            im.seek(0)
        if im.mode in ("RGBA", "LA") or (im.mode == "P" and "transparency" in im.info):
            out = BytesIO()
            im.convert("RGBA").save(out, format="PNG")
            return out.getvalue()
        else:
            out = BytesIO()
            im.convert("RGB").save(out, format="PNG")
            return out.getvalue()
    except Exception:
        return None


JS_COLLECT_TEXT = r"""
() => {
  const out = [];
  const norm = (s) => (s || "").replace(/\u00a0/g, " ").replace(/\s+/g, " ").trim();

  const visible = (el) => {
    try {
      const cs = getComputedStyle(el);
      if (cs.display === "none" || cs.visibility === "hidden" || cs.opacity === "0") return false;
      const r = el.getBoundingClientRect();
      if (r.width <= 0 || r.height <= 0) return false;
      return true;
    } catch { return true; }
  };

  const collectFromRoot = (root) => {
    if (!root) return;

    const elems = root.querySelectorAll ? root.querySelectorAll("h1,h2,h3,h4,p,li,blockquote,pre,span,div") : [];
    for (const el of elems) {
      if (!visible(el)) continue;
      const tag = (el.tagName || "").toLowerCase();
      if (["nav","header","footer","aside","button","input","textarea","select"].includes(tag)) continue;

      const t = norm(el.innerText || el.textContent || "");
      if (t && t.length >= 2 && t.length <= 4000) out.push(t);
    }

    const all = root.querySelectorAll ? root.querySelectorAll("*") : [];
    for (const el of all) {
      if (el.shadowRoot) collectFromRoot(el.shadowRoot);
    }
  };

  collectFromRoot(document);

  const bodyText = norm(document.body?.innerText || "");
  if (bodyText) out.push(bodyText);

  return out;
}
"""

JS_COLLECT_IMAGES = r"""
() => {
  const urls = [];
  const push = (u) => { if (u) urls.push(u); };

  const getBgUrls = (el) => {
    try {
      const cs = getComputedStyle(el);
      const bg = cs.backgroundImage || "";
      const matches = [...bg.matchAll(/url\(["']?(.*?)["']?\)/g)];
      return matches.map(m => m[1]).filter(Boolean);
    } catch { return []; }
  };

  const walk = (root) => {
    if (!root) return;
    if (root.shadowRoot) walk(root.shadowRoot);

    const nodes = root.querySelectorAll ? root.querySelectorAll("*") : [];
    for (const el of nodes) {
      const tag = (el.tagName || "").toLowerCase();
      if (tag === "img") {
        push(el.currentSrc || el.getAttribute("src") || el.getAttribute("data-src") || el.getAttribute("data-original") || el.getAttribute("data-lazy-src"));
      }
      for (const u of getBgUrls(el)) push(u);
      if (el.shadowRoot) walk(el);
    }
  };

  walk(document);
  return urls;
}
"""


async def wait_for_text_stable(page, interval_ms=700, stable_times=4, timeout_ms=45000):
    start = time.time()
    last_len = -1
    stable = 0
    while (time.time() - start) * 1000 < timeout_ms:
        try:
            length = await page.evaluate("() => (document.body?.innerText || '').length")
        except Exception:
            length = 0

        if length == last_len and length > 200:
            stable += 1
        else:
            stable = 0
        last_len = length

        if stable >= stable_times:
            return last_len

        await page.wait_for_timeout(interval_ms)
    return last_len


async def main():
    url = input("請輸入文章網址：\n").strip()
    if not url:
        print("未輸入網址，結束。")
        return

    # 用持久化 profile（反爬/渲染更穩）
    user_data_dir = os.path.join(os.getcwd(), "pw_profile_mcsm")

    async with async_playwright() as p:
        context = await p.chromium.launch_persistent_context(
            user_data_dir=user_data_dir,
            headless=False,
            viewport={"width": 1400, "height": 900},
            locale="zh-TW",
            args=["--disable-blink-features=AutomationControlled"],
        )
        page = await context.new_page()

        try:
            print("[INFO] Loading...")
            await page.goto(url, wait_until="domcontentloaded", timeout=60000)

            final_len = await wait_for_text_stable(page)
            print(f"[INFO] body_text_len(stable)={final_len}")

            # 觸發 lazy
            try:
                await page.evaluate("() => window.scrollTo(0, document.body.scrollHeight)")
                await page.wait_for_timeout(1200)
                await page.evaluate("() => window.scrollTo(0, 0)")
                await page.wait_for_timeout(600)
            except Exception:
                pass

            title = sanitize_text(await page.title()) or "Untitled"

            # 抓文字（含 shadow）
            lines = await page.evaluate(JS_COLLECT_TEXT)
            paras = normalize_paragraphs(lines)

            # 抓圖片（含 background-image）
            raw_imgs = await page.evaluate(JS_COLLECT_IMAGES)
            raw_imgs = uniq_keep_order(raw_imgs)

            img_urls = []
            for u in raw_imgs:
                try:
                    img_urls.append(urljoin(url, u))
                except Exception:
                    pass
            img_urls = uniq_keep_order(img_urls)

            print(f"[INFO] paras={len(paras)}, imgs={len(img_urls)}")

            if len(paras) < 3:
                print("[ERR] 正文仍不足（可能是 Canvas 文字）。若你要，我再給 OCR 版。")
                return

            # --- 產 docx（用 python-docx，Word 最穩） ---
            doc = Document()
            doc.add_heading(title, level=1)

            # 來源用 hxxps 避免 Word 當外部連結
            doc.add_paragraph(f"來源：{safe_source_url(url)}")
            doc.add_paragraph("")

            for ptxt in paras:
                ptxt = sanitize_text(ptxt)
                if ptxt:
                    doc.add_paragraph(ptxt)

            # 圖片（盡量轉 png 內嵌）
            # 若 PIL 不可用，則只內嵌 jpg/png，其他略過（避免壞檔/警告）
            added_any = False
            for u in img_urls:
                b, ct = fetch_image_bytes(u)
                if not b:
                    continue

                lower = u.lower()
                ct = (ct or "").lower()

                # svg 跳過（最容易出事）
                if "image/svg" in ct or lower.endswith(".svg"):
                    continue

                png = try_convert_to_png(b)
                if png is not None:
                    try:
                        doc.add_picture(BytesIO(png), width=Inches(6.0))
                        added_any = True
                    except Exception:
                        pass
                    continue

                # 沒 PIL 或轉失敗：只接受 jpg/png
                is_png = ("image/png" in ct) or lower.endswith(".png")
                is_jpg = ("image/jpeg" in ct) or ("image/jpg" in ct) or re.search(r"\.jpe?g(\?|#|$)", lower)
                if is_png or is_jpg:
                    try:
                        doc.add_picture(BytesIO(b), width=Inches(6.0))
                        added_any = True
                    except Exception:
                        pass

            # 檔名：用標題，避免覆蓋
            out_name = safe_filename(title) + ".docx"
            if os.path.exists(out_name):
                out_name = safe_filename(title) + "_" + datetime.now().strftime("%Y%m%d_%H%M%S") + ".docx"

            doc.save(out_name)
            print(f"[OK] 已輸出：{out_name}")
            if not added_any:
                print("[INFO] 圖片可能因格式（webp 等）或缺 PIL 而略過；文字已完成。")

        finally:
            await context.close()


if __name__ == "__main__":
    asyncio.run(main())
