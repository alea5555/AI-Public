import asyncio
import os
import re
import csv
import time
import traceback
from datetime import datetime
from io import BytesIO
from urllib.parse import urljoin, urlparse

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from docx.image.exceptions import UnrecognizedImageError

from playwright.async_api import async_playwright

# 可選：圖片轉 png（webp/gif 等更穩）
try:
    from PIL import Image
    PIL_OK = True
except Exception:
    PIL_OK = False


# ===== 你的路徑（沿用你批次習慣）=====
CSV_PATH  = r"F:\F\AI\web\web.csv"
OUT_DIR   = r"F:\F\AI\web"
SLEEP_SEC = 0.2  # 圖片下載間隔（可調）


# -----------------------------
# 檔名安全
# -----------------------------
def safe_filename(name: str, max_len: int = 120) -> str:
    name = (name or "").strip()
    if not name:
        return ""
    name = re.sub(r'[<>:"/\\|?*\x00-\x1F]', "_", name)
    name = re.sub(r"\s+", " ", name).strip()
    if len(name) > max_len:
        name = name[:max_len].rstrip()
    return name


def derive_name_from_url(url: str) -> str:
    try:
        u = urlparse(url)
        host = u.netloc or "site"
        path_last = (u.path or "").rstrip("/").split("/")[-1] or "index"
        q = (u.query or "").strip()
        if q:
            q = safe_filename(q)[:40]
            return safe_filename(f"{host}_{path_last}_{q}") or "site_index"
        return safe_filename(f"{host}_{path_last}") or "site_index"
    except Exception:
        return safe_filename(url)[:120] or "site_index"


# -----------------------------
# ✅ 清掉 Word 最討厭的控制字元（避免修復提示）
# -----------------------------
def sanitize_text(s: str) -> str:
    if s is None:
        return ""
    out_chars = []
    for ch in str(s):
        code = ord(ch)
        if code in (0x09, 0x0A, 0x0D):  # tab, lf, cr
            out_chars.append(ch)
            continue
        if code < 0x20:
            continue
        if 0xD800 <= code <= 0xDFFF:
            continue
        if code in (0xFFFE, 0xFFFF):
            continue
        out_chars.append(ch)

    out = "".join(out_chars)
    out = out.replace("\u00a0", " ")
    out = re.sub(r"\s+", " ", out).strip()
    return out


def uniq_keep_order(items):
    seen = set()
    out = []
    for x in items:
        x = (x or "").strip()
        if not x:
            continue
        if x in seen:
            continue
        seen.add(x)
        out.append(x)
    return out


def normalize_paragraphs(lines):
    # 去掉常見雜訊（需要可自行加）
    junk_contains = ["MULTISTRATEGYS", "策略市集", "策略生成器"]

    cleaned = []
    for s in lines:
        s = sanitize_text(s)
        if not s or len(s) < 2:
            continue
        if any(j in s for j in junk_contains):
            continue
        cleaned.append(s)

    cleaned = uniq_keep_order(cleaned)

    # 合併短碎片
    merged = []
    for s in cleaned:
        if not merged:
            merged.append(s)
            continue
        prev = merged[-1]
        if len(prev) < 12 and len(s) < 20 and not re.search(r"[。！？:：]$", prev):
            merged[-1] = sanitize_text(prev + s)
        else:
            merged.append(s)

    # 切段
    paras = []
    buf = []

    def flush():
        nonlocal buf
        t = sanitize_text(" ".join(buf))
        if t:
            paras.append(t)
        buf = []

    for s in merged:
        is_heading_like = bool(re.match(r"^([一二三四五六七八九十]、|[0-9]+[.)])", s)) or (len(s) <= 28 and "：" in s)
        if is_heading_like and buf:
            flush()

        buf.append(s)
        cur = " ".join(buf)
        if re.search(r"[。！？]$", s) or len(cur) > 260:
            flush()

    flush()
    return paras


# -----------------------------
# ✅ CSV 讀取（A:URL, B:名稱, C:前綴/序號，不解析）
# -----------------------------
def read_web_csv(csv_path: str):
    encodings_to_try = ["utf-8-sig", "cp950", "big5", "utf-8"]
    last_err = None

    for enc in encodings_to_try:
        try:
            rows = []
            with open(csv_path, "r", encoding=enc, newline="") as f:
                reader = csv.reader(f)
                for r in reader:
                    if not r:
                        continue
                    url = (r[0] if len(r) > 0 else "").strip()
                    if not url:
                        break
                    name_b = (r[1] if len(r) > 1 else "").strip()
                    prefix_c = (r[2] if len(r) > 2 else "").strip()  # ✅ C 欄原樣用（只做檔名安全化）
                    rows.append((url, name_b, prefix_c))
            return rows, enc
        except UnicodeDecodeError as e:
            last_err = e
            continue

    # 最後保底：cp950 replace
    try:
        rows = []
        with open(csv_path, "rb") as fb:
            raw = fb.read()
        text = raw.decode("cp950", errors="replace")
        for r in csv.reader(text.splitlines()):
            if not r:
                continue
            url = (r[0] if len(r) > 0 else "").strip()
            if not url:
                break
            name_b = (r[1] if len(r) > 1 else "").strip()
            prefix_c = (r[2] if len(r) > 2 else "").strip()
            rows.append((url, name_b, prefix_c))
        return rows, "cp950(replace)"
    except Exception:
        if last_err:
            raise last_err
        raise


# -----------------------------
# ✅ PO 文日期抽取（只在 C 欄空白時使用）
# -----------------------------
def _parse_date_to_yyyymmdd(s: str):
    if not s:
        return None
    s = s.strip()

    m = re.search(r"(20\d{2})-(\d{1,2})-(\d{1,2})", s)
    if m:
        y, mo, d = m.group(1), int(m.group(2)), int(m.group(3))
        return f"{y}{mo:02d}{d:02d}"

    m = re.search(r"(20\d{2})/(\d{1,2})/(\d{1,2})", s)
    if m:
        y, mo, d = m.group(1), int(m.group(2)), int(m.group(3))
        return f"{y}{mo:02d}{d:02d}"

    return None


def extract_date8_from_html(html: str) -> str:
    soup = BeautifulSoup(html, "lxml")

    meta_keys = [
        ("property", "article:published_time"),
        ("property", "og:published_time"),
        ("name", "pubdate"),
        ("name", "publishdate"),
        ("name", "publish_date"),
        ("name", "date"),
        ("itemprop", "datePublished"),
    ]

    for attr, val in meta_keys:
        tag = soup.find("meta", attrs={attr: val})
        if tag and tag.get("content"):
            d8 = _parse_date_to_yyyymmdd(tag["content"])
            if d8:
                return d8

    t = soup.find("time")
    if t:
        dt = t.get("datetime") or t.get_text(" ", strip=True)
        d8 = _parse_date_to_yyyymmdd(dt)
        if d8:
            return d8

    m = re.search(r"\b(20\d{2})[-/](\d{1,2})[-/](\d{1,2})\b", html)
    if m:
        y, mo, d = m.group(1), int(m.group(2)), int(m.group(3))
        return f"{y}{mo:02d}{d:02d}"

    return datetime.now().strftime("%Y%m%d")


# -----------------------------
# Playwright 抓正文/圖片（跟單次版同招）
# -----------------------------
JS_COLLECT_TEXT = r"""
() => {
  const out = [];
  const norm = (s) => (s || "").replace(/\u00a0/g, " ").replace(/\s+/g, " ").trim();

  const visible = (el) => {
    try {
      const cs = getComputedStyle(el);
      if (cs.display === "none" || cs.visibility === "hidden" || cs.opacity === "0") return false;
      const r = el.getBoundingClientRect();
      if (r.width <= 0 || r.height <= 0) return false;
      return true;
    } catch { return true; }
  };

  const collectFromRoot = (root) => {
    if (!root) return;

    const elems = root.querySelectorAll ? root.querySelectorAll("h1,h2,h3,h4,p,li,blockquote,pre,span,div") : [];
    for (const el of elems) {
      if (!visible(el)) continue;
      const tag = (el.tagName || "").toLowerCase();
      if (["nav","header","footer","aside","button","input","textarea","select"].includes(tag)) continue;

      const t = norm(el.innerText || el.textContent || "");
      if (t && t.length >= 2 && t.length <= 4000) out.push(t);
    }

    const all = root.querySelectorAll ? root.querySelectorAll("*") : [];
    for (const el of all) {
      if (el.shadowRoot) collectFromRoot(el.shadowRoot);
    }
  };

  collectFromRoot(document);

  const bodyText = norm(document.body?.innerText || "");
  if (bodyText) out.push(bodyText);

  return out;
}
"""

JS_COLLECT_IMAGES = r"""
() => {
  const urls = [];
  const push = (u) => { if (u) urls.push(u); };

  const getBgUrls = (el) => {
    try {
      const cs = getComputedStyle(el);
      const bg = cs.backgroundImage || "";
      const matches = [...bg.matchAll(/url\(["']?(.*?)["']?\)/g)];
      return matches.map(m => m[1]).filter(Boolean);
    } catch { return []; }
  };

  const walk = (root) => {
    if (!root) return;
    if (root.shadowRoot) walk(root.shadowRoot);

    const nodes = root.querySelectorAll ? root.querySelectorAll("*") : [];
    for (const el of nodes) {
      const tag = (el.tagName || "").toLowerCase();
      if (tag === "img") {
        push(el.currentSrc || el.getAttribute("src") || el.getAttribute("data-src") || el.getAttribute("data-original") || el.getAttribute("data-lazy-src"));
      }
      for (const u of getBgUrls(el)) push(u);
      if (el.shadowRoot) walk(el);
    }
  };

  walk(document);
  return urls;
}
"""


async def wait_for_text_stable(page, interval_ms=700, stable_times=4, timeout_ms=45000):
    start = time.time()
    last_len = -1
    stable = 0
    while (time.time() - start) * 1000 < timeout_ms:
        try:
            length = await page.evaluate("() => (document.body?.innerText || '').length")
        except Exception:
            length = 0

        if length == last_len and length > 200:
            stable += 1
        else:
            stable = 0
        last_len = length

        if stable >= stable_times:
            return last_len

        await page.wait_for_timeout(interval_ms)
    return last_len


def fetch_image_bytes(session: requests.Session, img_url: str, timeout: int = 30):
    try:
        r = session.get(img_url, timeout=timeout, headers={"User-Agent": "Mozilla/5.0"})
        if r.status_code != 200:
            return None, ""
        ct = (r.headers.get("content-type") or "").lower()
        if "image" not in ct:
            return None, ct
        return r.content, ct
    except Exception:
        return None, ""


def try_convert_to_png(image_bytes: bytes):
    if not PIL_OK:
        return None
    try:
        im = Image.open(BytesIO(image_bytes))
        if getattr(im, "is_animated", False):
            im.seek(0)
        out = BytesIO()
        if im.mode in ("RGBA", "LA") or (im.mode == "P" and "transparency" in im.info):
            im.convert("RGBA").save(out, format="PNG")
        else:
            im.convert("RGB").save(out, format="PNG")
        return out.getvalue()
    except Exception:
        return None


def safe_source_url(url: str) -> str:
    # 避免 Word 自動當外部連結觸發更新提示
    return url.replace("https://", "hxxps://").replace("http://", "hxxp://")


async def process_one(page, http_sess: requests.Session, url: str, name_b: str, prefix_c_raw: str, idx: int, total: int):
    url = (url or "").strip()
    if not url:
        return

    print(f"\n[{idx}/{total}] {url}")
    try:
        await page.goto(url, wait_until="domcontentloaded", timeout=60000)
    except Exception as e:
        print(f"❌ goto 失敗：{e}")
        return

    # 等正文穩定 + 觸發 lazy
    try:
        await wait_for_text_stable(page)
        await page.evaluate("() => window.scrollTo(0, document.body.scrollHeight)")
        await page.wait_for_timeout(1200)
        await page.evaluate("() => window.scrollTo(0, 0)")
        await page.wait_for_timeout(600)
    except Exception:
        pass

    # 抓渲染後 HTML（用來抽日期）
    try:
        html = await page.content()
    except Exception:
        html = ""

    # title（B 欄空白時使用）
    try:
        page_title = sanitize_text(await page.title())
    except Exception:
        page_title = ""

    # ✅ 前綴決策：C 欄有值就用「原樣」（只做檔名安全化）；C 空才用 PO 日期
    prefix_c = safe_filename(prefix_c_raw)
    if prefix_c:
        prefix = prefix_c
    else:
        prefix = extract_date8_from_html(html)

    # ✅ 名稱決策：B 欄 > title > url 推導
    file_base = safe_filename(name_b)
    if not file_base:
        file_base = safe_filename(page_title)
    if not file_base:
        file_base = derive_name_from_url(url)

    out_path = os.path.join(OUT_DIR, f"{prefix}_{file_base}.docx")

    if os.path.exists(out_path):
        print(f"⏭️ 已存在：{os.path.basename(out_path)}")
        return

    # ✅ 抓正文（含 shadow）
    try:
        lines = await page.evaluate(JS_COLLECT_TEXT)
    except Exception:
        lines = []
    paras = normalize_paragraphs(lines)

    # ✅ 抓圖片（含 background）
    try:
        raw_imgs = await page.evaluate(JS_COLLECT_IMAGES)
    except Exception:
        raw_imgs = []
    raw_imgs = uniq_keep_order(raw_imgs)

    img_urls = []
    for u in raw_imgs:
        try:
            img_urls.append(urljoin(url, u))
        except Exception:
            pass
    img_urls = uniq_keep_order(img_urls)

    print(f"[INFO] paras={len(paras)}, imgs={len(img_urls)}")

    # --- 產 docx ---
    doc = Document()
    doc.add_heading(page_title or file_base, level=1)
    doc.add_paragraph(f"來源：{safe_source_url(url)}")
    doc.add_paragraph(f"建檔前綴：{prefix}")
    doc.add_paragraph("")

    text_count = 0
    for ptxt in paras:
        ptxt = sanitize_text(ptxt)
        if ptxt:
            doc.add_paragraph(ptxt)
            text_count += 1

    img_count = 0
    for u in img_urls:
        lower = u.lower()
        if lower.endswith(".svg") or lower.endswith(".ico"):
            continue

        b, ct = fetch_image_bytes(http_sess, u)
        if not b:
            continue

        # svg 跳過
        if "image/svg" in (ct or ""):
            continue

        try:
            png = try_convert_to_png(b)
            if png is not None:
                doc.add_picture(BytesIO(png), width=Inches(6.0))
            else:
                # 沒 PIL：只收 jpg/png（避免 UnrecognizedImageError）
                is_png = ("image/png" in (ct or "")) or lower.endswith(".png")
                is_jpg = ("image/jpeg" in (ct or "")) or ("image/jpg" in (ct or "")) or re.search(r"\.jpe?g(\?|#|$)", lower)
                if not (is_png or is_jpg):
                    continue
                doc.add_picture(BytesIO(b), width=Inches(6.0))

            img_count += 1
            time.sleep(SLEEP_SEC)

        except UnrecognizedImageError:
            continue
        except Exception:
            continue

    doc.save(out_path)
    print(f"✅ 完成：{out_path}")
    print(f"📌 文字段落：約 {text_count} 段，圖片：{img_count} 張")


async def main_async():
    os.makedirs(OUT_DIR, exist_ok=True)

    items, enc = read_web_csv(CSV_PATH)
    print(f"✅ web.csv 讀取：{len(items)} 筆（encoding={enc}）")

    if not items:
        print("⚠️ web.csv 沒有網址（A欄）")
        return

    http_sess = requests.Session()
    http_sess.headers.update({"User-Agent": "Mozilla/5.0"})

    user_data_dir = os.path.join(os.getcwd(), "pw_profile_mcsm")

    async with async_playwright() as p:
        context = await p.chromium.launch_persistent_context(
            user_data_dir=user_data_dir,
            headless=False,
            viewport={"width": 1400, "height": 900},
            locale="zh-TW",
            args=["--disable-blink-features=AutomationControlled"],
        )
        page = await context.new_page()

        try:
            total = len(items)
            for idx, (url, name_b, prefix_c) in enumerate(items, start=1):
                try:
                    await process_one(page, http_sess, url, name_b, prefix_c, idx, total)
                except Exception as e:
                    print(f"❌ 這筆處理失敗：{e}")
                    continue
        finally:
            await context.close()


def main():
    try:
        asyncio.run(main_async())
    except Exception:
        print("❌ 程式發生未捕捉例外：")
        traceback.print_exc()


if __name__ == "__main__":
    main()
