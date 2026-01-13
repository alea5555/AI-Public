# -*- coding: utf-8 -*-
import re
import time
import traceback
from io import BytesIO
from pathlib import Path
from datetime import datetime
from urllib.parse import urlsplit, urlunsplit

from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Inches
from docx.image.exceptions import UnrecognizedImageError

# 可選：webp -> png（沒裝 PIL 也能跑，只是 webp 可能插不進 docx）
try:
    from PIL import Image
    PIL_OK = True
except Exception:
    PIL_OK = False


# =====================
# 你要的固定設定
# =====================
OUT_DIR = Path(r"F:\F\AI\FB")                 # ✅ DOCX 固定輸出到這裡
USER_DATA_DIR = r"F:\F\AI\pw_profile_fb"      # ✅ FB 登入持久化資料夾（第一次登入後會記住）
SLEEP_SEC = 0.25
MAX_IMAGES = 40


# =====================
# 工具
# =====================
def normalize_url(url: str) -> str:
    url = (url or "").strip()
    if not url:
        return ""
    if not url.startswith("http"):
        url = "https://" + url
    sp = urlsplit(url)
    return urlunsplit((sp.scheme, sp.netloc, sp.path, sp.query, ""))  # 去掉 #fragment


def safe_filename(s: str, max_len=120) -> str:
    s = re.sub(r'[<>:"/\\|?*]', "_", (s or "").strip())
    s = re.sub(r"\s+", "_", s)     # ✅ 空白 -> _
    s = re.sub(r"_+", "_", s)
    s = s.strip("_")
    if len(s) > max_len:
        s = s[:max_len].rstrip("_")
    return s or "Facebook"


def choose_available_path(folder: Path, base_name: str) -> Path:
    p0 = folder / f"{base_name}.docx"
    if not p0.exists():
        return p0
    for i in range(1, 200):
        p = folder / f"{base_name}_{i:02d}.docx"
        if not p.exists():
            return p
    return folder / f"{base_name}_{int(time.time())}.docx"


def date8_from_iso(dt_str: str) -> str:
    if not dt_str:
        return ""
    m = re.search(r"(20\d{2})-(\d{1,2})-(\d{1,2})", dt_str)
    if not m:
        return ""
    y, mo, d = m.group(1), int(m.group(2)), int(m.group(3))
    return f"{y}{mo:02d}{d:02d}"


def clean_text(raw: str) -> str:
    raw = raw or ""
    lines = []
    for ln in raw.splitlines():
        ln = ln.strip()
        if not ln:
            continue
        # 常見 UI 噪音（短字）
        bad_short = {"讚", "留言", "分享", "最相關", "更多", "查看更多", "查看翻譯", "回覆", "已編輯"}
        if ln in bad_short:
            continue
        lines.append(ln)

    # 去重
    out = []
    seen = set()
    for ln in lines:
        if ln in seen:
            continue
        seen.add(ln)
        out.append(ln)
    return "\n".join(out).strip()


# =====================
# 選最佳容器：dialog > article > main
# =====================
def get_best_container(page):
    dialogs = page.locator('div[role="dialog"]')
    best = None
    best_score = -1

    for i in range(dialogs.count()):
        d = dialogs.nth(i)
        aria = (d.get_attribute("aria-label") or "").lower()
        if "messenger" in aria or "chat" in aria:
            continue

        score = 0
        try:
            if d.locator('div[data-ad-preview="message"]').count() > 0:
                score += 8
        except Exception:
            pass
        try:
            if d.locator("time[datetime]").count() > 0:
                score += 7
        except Exception:
            pass
        try:
            score += min(d.locator('div[dir="auto"]').count(), 10)
        except Exception:
            pass

        if score > best_score:
            best_score = score
            best = d

    if best is not None and best_score >= 7:
        return best

    art = page.locator('div[role="article"]').first
    if art.count():
        return art

    main = page.locator('div[role="main"]').first
    if main.count():
        return main

    return page.locator("body").first


# =====================
# 抓 PO文日期（最重要）
# =====================
def extract_post_datetime(container, page) -> str:
    """
    ✅ 你要的 PO 文日期來源：
    優先抓貼文區塊內的 time[datetime]（最準）
    """
    # 1) container 內 time[datetime]
    try:
        t = container.locator("time[datetime]").first
        if t.count():
            dt = (t.get_attribute("datetime") or "").strip()
            if dt:
                return dt
    except Exception:
        pass

    # 2) 全頁 time[datetime]（退路）
    try:
        t = page.locator("time[datetime]").first
        if t.count():
            dt = (t.get_attribute("datetime") or "").strip()
            if dt:
                return dt
    except Exception:
        pass

    return ""


def extract_title(page) -> str:
    try:
        t = page.locator('meta[property="og:title"]').get_attribute("content")
        if t:
            return t.strip()
    except Exception:
        pass
    try:
        t = page.title()
        if t:
            return t.strip()
    except Exception:
        pass
    return "Facebook"


def extract_text(container) -> str:
    # 1) data-ad-preview="message"
    try:
        m = container.locator('div[data-ad-preview="message"]').first
        if m.count():
            return clean_text(m.inner_text(timeout=8000))
    except Exception:
        pass

    # 2) dir=auto
    try:
        d = container.locator('div[dir="auto"]').first
        if d.count():
            return clean_text(d.inner_text(timeout=8000))
    except Exception:
        pass

    # 3) container 全部文字
    try:
        return clean_text(container.inner_text(timeout=8000))
    except Exception:
        return ""


def collect_images_bytes(context_request, container):
    imgs_bytes = []
    seen = set()

    loc = container.locator('img[data-visualcompletion="media-vc-image"]')
    if loc.count() == 0:
        loc = container.locator("img")

    for i in range(min(loc.count(), MAX_IMAGES)):
        im = loc.nth(i)
        try:
            src = (im.get_attribute("src") or "").strip()
            if not src or src in seen:
                continue
            seen.add(src)

            # 過濾非常小的 icon/頭像（寬高取不到就不擋）
            w = im.get_attribute("width")
            h = im.get_attribute("height")
            try:
                wi = int(w) if (w and str(w).isdigit()) else 999
                hi = int(h) if (h and str(h).isdigit()) else 999
                if wi < 80 or hi < 80:
                    continue
            except Exception:
                pass

            r = context_request.get(src, timeout=25000)
            if not r.ok:
                continue
            b = r.body()
            if b:
                imgs_bytes.append(b)
        except Exception:
            continue

    return imgs_bytes


def add_image_to_doc(doc: Document, b: bytes) -> bool:
    if not b:
        return False

    # 先嘗試用 PIL 轉 PNG
    if PIL_OK:
        try:
            im = Image.open(BytesIO(b))
            out = BytesIO()
            im.convert("RGB").save(out, format="PNG")
            out.seek(0)
            doc.add_picture(out, width=Inches(6.3))
            return True
        except Exception:
            pass

    try:
        doc.add_picture(BytesIO(b), width=Inches(6.3))
        return True
    except UnrecognizedImageError:
        return False
    except Exception:
        return False


# =====================
# MAIN
# =====================
def main():
    url = normalize_url(input("請輸入 FB 貼文網址：\n").strip())
    if not url:
        print("❌ 未輸入網址，結束")
        return

    OUT_DIR.mkdir(parents=True, exist_ok=True)

    with sync_playwright() as p:
        # ✅ 重點：持久化 context（保留登入）
        context = p.chromium.launch_persistent_context(
            user_data_dir=USER_DATA_DIR,
            headless=False,       # FB 建議有頭模式
            locale="zh-TW",
            viewport={"width": 1366, "height": 900},
        )
        page = context.new_page()

        page.goto(url, wait_until="domcontentloaded", timeout=60000)
        page.wait_for_timeout(3000)

        container = get_best_container(page)

        title = extract_title(page)

        post_dt_iso = extract_post_datetime(container, page)
        post_date8 = date8_from_iso(post_dt_iso)

        # ✅ 抓不到就警告（不再默默用今天）
        if not post_date8:
            print("⚠️ 警告：未抓到 PO文日期（time[datetime]），可能尚未登入或貼文未完整載入。")
            post_date8 = datetime.now().strftime("%Y%m%d")

        content = extract_text(container)
        images = collect_images_bytes(context.request, container)

        # ✅ 關閉持久化 context（但登入會保留在 USER_DATA_DIR）
        context.close()

    # ✅ 檔名：PO文日期 + 標題
    base = safe_filename(title)
    out_path = choose_available_path(OUT_DIR, f"{post_date8}_{base}")

    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"來源網址：{url}")
    doc.add_paragraph(f"PO文日期：{post_date8}")
    if post_dt_iso:
        doc.add_paragraph(f"PO文時間(datetime)：{post_dt_iso}")
    doc.add_paragraph("")

    if content:
        for line in content.splitlines():
            doc.add_paragraph(line)
    else:
        doc.add_paragraph("（未成功抽取到正文，可能需要登入或貼文權限受限）")

    img_ok = 0
    if images:
        doc.add_paragraph("")
        doc.add_paragraph("【圖片】")
        for b in images[:30]:
            if add_image_to_doc(doc, b):
                img_ok += 1
                time.sleep(SLEEP_SEC)

    doc.save(out_path)

    print(f"✅ 完成：{out_path}")
    print(f"📌 圖片：{img_ok} 張")
    print(f"📅 PO文日期(YYYYMMDD)：{post_date8}")


if __name__ == "__main__":
    try:
        main()
    except Exception:
        print("❌ 程式發生未處理例外：")
        traceback.print_exc()
        # ✅ 不暫停，直接回 CMD
