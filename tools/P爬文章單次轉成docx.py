import os
import re
import time
import json
from io import BytesIO
from urllib.parse import urljoin, urlparse
from datetime import datetime

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from docx.image.exceptions import UnrecognizedImageError

# 可選：用來把 webp 轉 png（沒裝也沒關係，會自動跳過）
try:
    from PIL import Image
    PIL_OK = True
except Exception:
    PIL_OK = False


OUT_DIR = r"F:\F\AI"
SLEEP_SEC = 0.5  # 下載圖片間隔


def safe_filename(name: str, max_len: int = 120) -> str:
    name = re.sub(r'[<>:"/\\|?*]', "_", name).strip()
    name = re.sub(r"\s+", " ", name)
    if len(name) > max_len:
        name = name[:max_len].rstrip()
    return name or "output"


def fetch_html(session: requests.Session, url: str) -> str:
    url = url.split("#", 1)[0]
    r = session.get(url, timeout=30)
    r.raise_for_status()
    if not r.encoding or r.encoding.lower() == "iso-8859-1":
        r.encoding = r.apparent_encoding or "utf-8"
    return r.text


# =========================
# ✅ 改良：選「最像正文」的容器（不再只取第一個命中）
# =========================
def _node_score(node) -> int:
    if not node:
        return -10**9

    txt = node.get_text(" ", strip=True)
    tlen = len(txt)

    p = len(node.find_all("p"))
    li = len(node.find_all("li"))
    h = len(node.find_all(["h1", "h2", "h3", "h4"]))
    pre = len(node.find_all("pre"))
    code = len(node.find_all("code"))
    bq = len(node.find_all("blockquote"))
    img = len(node.find_all("img"))

    bad = 0
    for bad_sel in ["nav", "header", "footer", "aside"]:
        bad += len(node.find_all(bad_sel))

    cls = " ".join(node.get("class", [])).lower()
    nid = (node.get("id") or "").lower()
    if any(k in cls for k in ["comment", "sidebar", "related", "recommend", "widget", "breadcrumb", "footer"]):
        bad += 10
    if any(k in nid for k in ["comment", "sidebar", "related", "recommend", "footer"]):
        bad += 10

    score = 0
    score += min(tlen, 20000)
    score += p * 300
    score += li * 120
    score += h * 200
    score += pre * 200
    score += code * 50
    score += bq * 150
    score += img * 10
    score -= bad * 500
    return score


def pick_content_root(soup: BeautifulSoup):
    selectors = [
        "article",
        "main",
        ".vditor-reset",
        ".markdown-body",
        ".post-detail",
        ".post-content",
        ".entry-content",
        ".article-content",
        ".content",
        "#content",
        "#__next",
        "body",
    ]

    candidates = []
    for sel in selectors:
        for node in soup.select(sel):
            candidates.append(node)

    if not candidates:
        return soup.body or soup

    best = max(candidates, key=_node_score)
    return best


def is_probably_nav_or_junk(tag) -> bool:
    if tag.name in {"nav", "header", "footer", "aside", "script", "style", "noscript"}:
        return True
    cls = " ".join(tag.get("class", [])).lower()
    if any(k in cls for k in ["share", "related", "sidebar", "widget", "comment", "ads", "advert", "breadcrumb"]):
        return True
    return False


def iter_content_blocks(root):
    for t in root.find_all(["script", "style", "noscript"]):
        t.decompose()

    for el in root.find_all(["h1", "h2", "h3", "h4", "p", "li", "blockquote", "pre", "img"]):
        if is_probably_nav_or_junk(el):
            continue

        if el.name in ["h1", "h2", "h3", "h4"]:
            txt = el.get_text(" ", strip=True)
            if txt:
                yield ("heading", el.name, txt)
            continue

        if el.name == "p":
            txt = el.get_text(" ", strip=True)
            if txt:
                yield ("p", txt)
            continue

        if el.name == "li":
            txt = el.get_text(" ", strip=True)
            if txt:
                yield ("li", txt)
            continue

        if el.name == "blockquote":
            txt = el.get_text(" ", strip=True)
            if txt:
                yield ("quote", txt)
            continue

        if el.name == "pre":
            txt = el.get_text("\n", strip=True)
            if txt:
                yield ("codeblock", txt)
            continue

        if el.name == "img":
            src = el.get("src") or el.get("data-src") or el.get("data-lazy-src") or el.get("data-original")
            if not src:
                continue
            alt = (el.get("alt") or "").strip()
            yield ("img", src, alt)
            continue


def download_image(session: requests.Session, img_url: str):
    try:
        r = session.get(img_url, timeout=30)
        r.raise_for_status()
        ctype = (r.headers.get("Content-Type") or "").lower()
        # 不是圖片就不要
        if "image" not in ctype:
            return None, ctype
        return r.content, ctype
    except Exception:
        return None, ""


# --------- 保底：從 Next/Nuxt JSON 找正文 ---------
def _extract_next_data_json(html: str):
    m = re.search(r'<script[^>]+id="__NEXT_DATA__"[^>]*>(.*?)</script>', html, re.S | re.I)
    if not m:
        return None
    try:
        return json.loads(m.group(1).strip())
    except Exception:
        return None


def _extract_nuxt_data_json(html: str):
    m = re.search(r'window\.__NUXT__\s*=\s*(\{.*?\});\s*</script>', html, re.S | re.I)
    if not m:
        m = re.search(r'window\.__NUXT__\s*=\s*(\{.*\})\s*;?', html, re.S | re.I)
    if not m:
        return None
    raw = m.group(1).strip()
    raw = re.sub(r"\bundefined\b", "null", raw)
    try:
        return json.loads(raw)
    except Exception:
        return None


def _score_candidate(key: str, s: str) -> int:
    k = (key or "").lower()
    score = 0
    if any(x in k for x in ["content", "html", "markdown", "body", "article", "text"]):
        score += 200
    if "<p" in s or "<h" in s or "</" in s:
        score += 150
    if "\n" in s:
        score += 50
    if "function(" in s or "var " in s or "webpack" in s:
        score -= 300
    score += min(len(s) // 50, 400)
    return score


def _find_best_long_text(obj, key_path=""):
    best = ("", "", -10**9)

    def walk(x, kp):
        nonlocal best
        if isinstance(x, dict):
            for k, v in x.items():
                walk(v, f"{kp}.{k}" if kp else str(k))
        elif isinstance(x, list):
            for i, v in enumerate(x):
                walk(v, f"{kp}[{i}]")
        elif isinstance(x, str):
            s = x.strip()
            if len(s) < 200:
                return
            sc = _score_candidate(kp.split(".")[-1], s)
            if sc > best[2]:
                best = (kp, s, sc)

    walk(obj, key_path)
    return best


def try_extract_article_text_from_scripts(html: str):
    data = _extract_next_data_json(html)
    if not data:
        data = _extract_nuxt_data_json(html)
    if not data:
        return None

    kp, text, score = _find_best_long_text(data)
    if score < 0:
        return None
    return text


def add_plaintext_to_doc(doc: Document, text: str):
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    blocks = [b.strip() for b in re.split(r"\n{2,}", text) if b.strip()]
    for b in blocks:
        if re.match(r"^#{1,4}\s+", b):
            level = len(re.match(r"^(#+)", b).group(1))
            title = re.sub(r"^#{1,4}\s+", "", b).strip()
            doc.add_heading(title, level=min(level, 4))
        else:
            lines = b.split("\n")
            p = doc.add_paragraph(lines[0])
            for line in lines[1:]:
                p.add_run("\n" + line)


# =========================
# ✅ 新增：抓「建檔日期」→ yyyymmdd（抓不到就用今天）
# =========================
def extract_date8(soup: BeautifulSoup, html: str) -> str:
    # 常見 meta
    meta_keys = [
        ("property", "article:published_time"),
        ("property", "og:published_time"),
        ("name", "pubdate"),
        ("name", "publishdate"),
        ("name", "publish_date"),
        ("name", "date"),
        ("itemprop", "datePublished"),
    ]

    for attr, val in meta_keys:
        tag = soup.find("meta", attrs={attr: val})
        if tag and tag.get("content"):
            d8 = _parse_date_to_yyyymmdd(tag["content"])
            if d8:
                return d8

    # <time datetime="...">
    t = soup.find("time")
    if t:
        dt = t.get("datetime") or t.get_text(" ", strip=True)
        d8 = _parse_date_to_yyyymmdd(dt)
        if d8:
            return d8

    # HTML 裡直接找 yyyy-mm-dd 或 yyyy/mm/dd
    m = re.search(r"\b(20\d{2})[-/](\d{1,2})[-/](\d{1,2})\b", html)
    if m:
        y, mo, d = m.group(1), int(m.group(2)), int(m.group(3))
        return f"{y}{mo:02d}{d:02d}"

    # 抓不到 → 今天
    return datetime.now().strftime("%Y%m%d")


def _parse_date_to_yyyymmdd(s: str) -> str | None:
    if not s:
        return None
    s = s.strip()

    # 2025-03-28T... / 2025-03-28
    m = re.search(r"(20\d{2})-(\d{1,2})-(\d{1,2})", s)
    if m:
        y, mo, d = m.group(1), int(m.group(2)), int(m.group(3))
        return f"{y}{mo:02d}{d:02d}"

    # 2025/03/28
    m = re.search(r"(20\d{2})/(\d{1,2})/(\d{1,2})", s)
    if m:
        y, mo, d = m.group(1), int(m.group(2)), int(m.group(3))
        return f"{y}{mo:02d}{d:02d}"

    return None


# =========================
# ✅ 新增：把 webp 轉 png（可選）
# =========================
def maybe_convert_webp_to_png_bytes(img_bytes: bytes, ctype: str, img_url: str) -> bytes | None:
    low_ct = (ctype or "").lower()
    ext = os.path.splitext(urlparse(img_url).path.lower())[1]

    is_webp = ("image/webp" in low_ct) or (ext == ".webp")
    if not is_webp:
        return None

    if not PIL_OK:
        return None

    try:
        im = Image.open(BytesIO(img_bytes))
        out = BytesIO()
        im.convert("RGB").save(out, format="PNG")
        return out.getvalue()
    except Exception:
        return None


def main():
    url = input("請輸入文章網址：\n").strip()
    if not url:
        print("❌ 未輸入網址，結束")
        return

    os.makedirs(OUT_DIR, exist_ok=True)

    with requests.Session() as s:
        s.headers.update({
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8",
            "Accept-Language": "zh-TW,zh;q=0.9,en;q=0.6",
            "Referer": "https://www.codefather.cn/",
        })

        html = fetch_html(s, url)
        soup = BeautifulSoup(html, "lxml")

        page_title = soup.title.get_text(strip=True) if soup.title else "article"
        file_base = safe_filename(page_title)

        # ✅ 只加這個：yyyymmdd_ 前綴
        date8 = extract_date8(soup, html)
        out_path = os.path.join(OUT_DIR, f"{date8}_{file_base}.docx")

        doc = Document()
        doc.add_heading(page_title, level=0)
        doc.add_paragraph(f"來源網址：{url}")
        doc.add_paragraph(f"建檔日期：{date8}")

        root = pick_content_root(soup)

        img_count = 0
        text_count = 0

        blocks = list(iter_content_blocks(root))
        for block in blocks:
            kind = block[0]

            if kind == "heading":
                _, tagname, txt = block
                level_map = {"h1": 1, "h2": 2, "h3": 3, "h4": 4}
                doc.add_heading(txt, level=level_map.get(tagname, 2))
                text_count += 1

            elif kind == "p":
                _, txt = block
                doc.add_paragraph(txt)
                text_count += 1

            elif kind == "li":
                _, txt = block
                doc.add_paragraph(txt, style="List Bullet")
                text_count += 1

            elif kind == "quote":
                _, txt = block
                doc.add_paragraph(txt, style="Intense Quote")
                text_count += 1

            elif kind == "codeblock":
                _, txt = block
                p = doc.add_paragraph()
                run = p.add_run(txt)
                run.font.name = "Consolas"
                text_count += 1

            elif kind == "img":
                _, src, alt = block
                img_url = urljoin(url.split("#", 1)[0], src)

                # 先用副檔名判斷明顯不支援的（svg/ico）
                path = urlparse(img_url).path.lower()
                if any(path.endswith(x) for x in [".svg", ".ico"]):
                    continue

                img, ctype = download_image(s, img_url)
                if not img:
                    continue

                if alt:
                    doc.add_paragraph(alt)

                # ✅ 這段是關鍵：不再因為圖片不認得而整支崩掉
                try:
                    # 若是 webp，嘗試轉 png 再插
                    converted = maybe_convert_webp_to_png_bytes(img, ctype, img_url)
                    if converted:
                        doc.add_picture(BytesIO(converted), width=Inches(6.0))
                    else:
                        doc.add_picture(BytesIO(img), width=Inches(6.0))

                    img_count += 1
                    time.sleep(SLEEP_SEC)

                except UnrecognizedImageError:
                    print(f"⚠️ 無法識別圖片格式，跳過：{img_url}")
                    continue
                except Exception as e:
                    print(f"⚠️ 圖片插入失敗，跳過：{img_url} | {e}")
                    continue

        # ✅ 保底：如果 DOM 幾乎抓不到文字，就從 script JSON 抽正文
        if text_count <= 2:
            extracted = try_extract_article_text_from_scripts(html)
            if extracted:
                doc.add_page_break()
                doc.add_heading("（保底抽取內容）", level=1)

                if "<p" in extracted or "<h" in extracted or "</" in extracted:
                    soup2 = BeautifulSoup(extracted, "lxml")
                    root2 = pick_content_root(soup2)
                    for block in iter_content_blocks(root2):
                        kind = block[0]
                        if kind == "heading":
                            _, tagname, txt = block
                            level_map = {"h1": 1, "h2": 2, "h3": 3, "h4": 4}
                            doc.add_heading(txt, level=level_map.get(tagname, 2))
                            text_count += 1
                        elif kind == "p":
                            _, txt = block
                            doc.add_paragraph(txt)
                            text_count += 1
                        elif kind == "li":
                            _, txt = block
                            doc.add_paragraph(txt, style="List Bullet")
                            text_count += 1
                        elif kind == "quote":
                            _, txt = block
                            doc.add_paragraph(txt, style="Intense Quote")
                            text_count += 1
                        elif kind == "codeblock":
                            _, txt = block
                            p = doc.add_paragraph()
                            run = p.add_run(txt)
                            run.font.name = "Consolas"
                            text_count += 1
                else:
                    add_plaintext_to_doc(doc, extracted)
                    text_count += 1

        doc.save(out_path)
        print(f"✅ 完成：{out_path}")
        print(f"📌 文字段落/項目：約 {text_count} 份，圖片：{img_count} 張")


if __name__ == "__main__":
    main()
