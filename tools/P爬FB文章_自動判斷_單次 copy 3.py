# -*- coding: utf-8 -*-
import os
import re
import time
import traceback
from io import BytesIO
from urllib.parse import urlparse
from datetime import datetime

import requests
from docx import Document
from docx.shared import Inches
from docx.image.exceptions import UnrecognizedImageError

# 可選：webp -> png（沒裝 PIL 也能跑，只是 webp 可能跳過）
try:
    from PIL import Image
    PIL_OK = True
except Exception:
    PIL_OK = False

OUT_DIR = r"F:\F\AI"
SLEEP_SEC = 0.35
MAX_MEDIA = 40


def safe_filename(name: str, max_len: int = 120) -> str:
    name = re.sub(r'[<>:"/\\|?*]', "_", str(name)).strip()
    name = re.sub(r"\s+", "_", name)          # ✅ 空白 -> _
    name = re.sub(r"_+", "_", name)           # 合併多個 _
    if len(name) > max_len:
        name = name[:max_len].rstrip("_")
    return name or "output"


def _browser_headers(referer: str = "https://www.threads.com/") -> dict:
    return {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/122.0.0.0 Safari/537.36"
        ),
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8",
        "Accept-Language": "zh-TW,zh;q=0.9,en;q=0.8",
        "Referer": referer,
        "Connection": "keep-alive",
        "Upgrade-Insecure-Requests": "1",
    }


def parse_threads_url_info(url: str) -> str:
    u = urlparse(url)
    parts = [p for p in u.path.split("/") if p]
    handle = ""
    post_id = ""
    for i, p in enumerate(parts):
        if p.startswith("@"):
            handle = p.lstrip("@")
        if p == "post" and i + 1 < len(parts):
            post_id = parts[i + 1]
    if handle and post_id:
        return f"threads_{handle}_{post_id}"
    if post_id:
        return f"threads_{post_id}"
    return "threads_post"


def date8_from_any_datetime(dt_str: str) -> str:
    """
    從 Threads DOM time[datetime] 的 ISO 轉 YYYYMMDD
    e.g. 2026-01-07T13:22:00.000Z -> 20260107
    """
    if not dt_str:
        return ""
    m = re.search(r"(20\d{2})-(\d{1,2})-(\d{1,2})", dt_str)
    if not m:
        return ""
    y, mo, d = m.group(1), int(m.group(2)), int(m.group(3))
    return f"{y}{mo:02d}{d:02d}"


def fetch_playwright_bundle(url: str):
    """
    ✅ 用 Playwright 抓：
    - PO文時間：從 time[datetime]（你要的「網頁上的日期來源」）
    - 可視 DOM 正文：找 main 內最長可視文字區塊
    - 貼文圖片：main 內 img 的 currentSrc/src/srcset，過濾小圖
    - 整頁截圖：抓不到圖時保底
    - 標題：og:title（沒有就用網址推導）
    """
    from playwright.sync_api import sync_playwright

    headless = os.environ.get("HEADLESS", "1") != "0"

    with sync_playwright() as p:
        browser = p.chromium.launch(
            headless=headless,
            args=[
                "--no-sandbox",
                "--disable-setuid-sandbox",
                "--disable-dev-shm-usage",
                "--disable-blink-features=AutomationControlled",
            ],
        )
        context = browser.new_context(
            user_agent=_browser_headers()["User-Agent"],
            locale="zh-TW",
            viewport={"width": 1366, "height": 900},
        )
        page = context.new_page()
        page.goto(url, wait_until="domcontentloaded", timeout=60000)

        # 等內容出現 + lazy-load
        page.wait_for_timeout(2200)
        try:
            page.mouse.wheel(0, 1800)
            page.wait_for_timeout(1200)
        except Exception:
            pass

        # 1) ✅ PO文時間（網頁 DOM 的 time[datetime]）
        post_datetime = ""
        try:
            post_datetime = page.evaluate(
                """() => {
                    const main = document.querySelector('main') || document.body;
                    const t1 = main.querySelector('time[datetime]');
                    if (t1 && t1.getAttribute('datetime')) return t1.getAttribute('datetime');

                    const t2 = document.querySelector('time[datetime]');
                    if (t2 && t2.getAttribute('datetime')) return t2.getAttribute('datetime');

                    return '';
                }"""
            ) or ""
            post_datetime = post_datetime.strip()
        except Exception:
            post_datetime = ""

        # 2) 正文（可視 DOM：最長文字區塊）
        post_text = ""
        try:
            post_text = page.evaluate(
                """() => {
                    const norm = (s) => (s || '')
                        .replace(/\\r/g, '')
                        .replace(/\\n{3,}/g, '\\n\\n')
                        .trim();

                    const isBad = (el) => {
                        if (!el) return true;
                        const tag = (el.tagName || '').toLowerCase();
                        if (['script','style','noscript','svg'].includes(tag)) return true;
                        const role = (el.getAttribute && el.getAttribute('role')) || '';
                        if (role && ['navigation','banner','dialog'].includes(role)) return true;
                        return false;
                    };

                    const main = document.querySelector('main') || document.body;
                    const cand = [];
                    const walker = document.createTreeWalker(main, NodeFilter.SHOW_ELEMENT, null);

                    while (walker.nextNode()) {
                        const el = walker.currentNode;
                        if (isBad(el)) continue;

                        const tag = (el.tagName || '').toLowerCase();
                        if (['button','input','textarea','select'].includes(tag)) continue;

                        const st = window.getComputedStyle(el);
                        if (!st || st.display === 'none' || st.visibility === 'hidden') continue;

                        const txt = norm(el.innerText);
                        if (!txt || txt.length < 80) continue;

                        const score = txt.length + (txt.includes('\\n') ? 120 : 0);
                        cand.push({score, txt});
                    }

                    cand.sort((a,b) => b.score - a.score);
                    if (!cand.length) return '';

                    const blacklist = ['Meta.ai', 'Cookie', '登入', 'Log in', 'Sign up'];
                    for (const c of cand) {
                        if (!blacklist.some(k => c.txt.includes(k))) return c.txt;
                    }
                    return cand[0].txt;
                }"""
            ) or ""
            post_text = post_text.strip()
        except Exception:
            post_text = ""

        # 3) 圖片（main 內大圖 currentSrc/srcset）
        img_urls = []
        try:
            img_urls = page.evaluate(
                """() => {
                    const out = new Set();
                    const main = document.querySelector('main') || document.body;
                    const imgs = Array.from(main.querySelectorAll('img'));

                    for (const img of imgs) {
                        const w = img.naturalWidth || 0;
                        const h = img.naturalHeight || 0;
                        // 過濾小圖（頭像/ICON）
                        if (w && h && (w < 120 || h < 120)) continue;

                        if (img.currentSrc) out.add(img.currentSrc);
                        if (img.src) out.add(img.src);

                        const ss = img.getAttribute('srcset');
                        if (ss) {
                            ss.split(',')
                              .map(s => s.trim().split(' ')[0])
                              .forEach(u => { if (u) out.add(u); });
                        }
                    }

                    // meta 圖（有時是縮圖）
                    const og = document.querySelector('meta[property="og:image"]');
                    if (og && og.content) out.add(og.content);

                    return Array.from(out);
                }"""
            ) or []
        except Exception:
            img_urls = []

        # 4) 截圖保底
        screenshot_bytes = None
        try:
            screenshot_bytes = page.screenshot(full_page=True)
        except Exception:
            screenshot_bytes = None

        # 5) 標題
        meta_title = ""
        try:
            meta_title = page.locator('meta[property="og:title"]').get_attribute("content") or ""
            meta_title = meta_title.strip()
        except Exception:
            meta_title = ""

        browser.close()

    # 清理圖片URL
    cleaned = []
    seen = set()
    for u in (img_urls or []):
        u = (u or "").strip()
        if not u or u.startswith("data:"):
            continue
        pth = urlparse(u).path.lower()
        if pth.endswith(".svg") or pth.endswith(".ico"):
            continue
        if u in seen:
            continue
        seen.add(u)
        cleaned.append(u)

    return meta_title, post_datetime, post_text, cleaned, screenshot_bytes


def download_image_bytes(session: requests.Session, img_url: str):
    try:
        r = session.get(img_url, timeout=30)
        r.raise_for_status()
        ctype = (r.headers.get("Content-Type") or "").lower()
        if "image" not in ctype:
            return None, ctype
        return r.content, ctype
    except Exception:
        return None, ""


def maybe_convert_webp_to_png_bytes(img_bytes: bytes, ctype: str, img_url: str):
    low_ct = (ctype or "").lower()
    ext = os.path.splitext(urlparse(img_url).path.lower())[1]
    is_webp = ("image/webp" in low_ct) or (ext == ".webp")
    if not is_webp or not PIL_OK:
        return None
    try:
        im = Image.open(BytesIO(img_bytes))
        out = BytesIO()
        im.convert("RGB").save(out, format="PNG")
        return out.getvalue()
    except Exception:
        return None


def add_picture_to_doc(doc: Document, img_bytes: bytes, width_inches: float = 6.3) -> bool:
    try:
        doc.add_picture(BytesIO(img_bytes), width=Inches(width_inches))
        return True
    except UnrecognizedImageError:
        return False
    except Exception:
        return False


def main():
    url = input("請輸入 Threads 貼文網址：\n").strip()
    if not url:
        print("❌ 未輸入網址，結束")
        return

    os.makedirs(OUT_DIR, exist_ok=True)

    meta_title, post_datetime, dom_text, img_urls, screenshot_bytes = fetch_playwright_bundle(url)

    # ✅ 你要的：檔名日期 = PO文日期（來自 time[datetime]）
    post_date8 = date8_from_any_datetime(post_datetime)
    if not post_date8:
        # 抓不到時才退回今天（極少數情況：貼文受限）
        post_date8 = datetime.now().strftime("%Y%m%d")

    # 標題
    fallback_title = parse_threads_url_info(url)
    title = meta_title if meta_title and meta_title.lower() != "threads" else fallback_title

    # 檔名
    out_path = os.path.join(OUT_DIR, f"{post_date8}_{safe_filename(title)}.docx")

    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"來源網址：{url}")
    doc.add_paragraph(f"PO文日期：{post_date8}")
    if post_datetime:
        doc.add_paragraph(f"PO文時間(datetime)：{post_datetime}")
    doc.add_paragraph("")

    if dom_text:
        doc.add_paragraph(dom_text)
    else:
        doc.add_paragraph("（未成功抽取到正文，可能貼文權限受限或需登入）")

    # 插入圖片
    img_urls = (img_urls or [])[:MAX_MEDIA]
    img_count = 0

    with requests.Session() as s:
        s.headers.update(_browser_headers(referer="https://www.threads.com/"))

        for img_url in img_urls:
            img, ctype = download_image_bytes(s, img_url)
            if not img:
                continue

            converted = maybe_convert_webp_to_png_bytes(img, ctype, img_url)
            ok = False
            if converted:
                ok = add_picture_to_doc(doc, converted)
            if not ok:
                ok = add_picture_to_doc(doc, img)

            if ok:
                img_count += 1
                time.sleep(SLEEP_SEC)

    # 沒圖片就插截圖保底
    if img_count == 0 and screenshot_bytes:
        doc.add_page_break()
        doc.add_heading("貼文截圖", level=1)
        if add_picture_to_doc(doc, screenshot_bytes, width_inches=6.8):
            img_count = 1

    doc.save(out_path)
    print(f"✅ 完成：{out_path}")
    print(f"📌 圖片：{img_count} 張")
    if post_datetime:
        print(f"🕒 PO文時間(datetime)：{post_datetime}")
    print(f"📅 PO文日期(YYYYMMDD)：{post_date8}")


if __name__ == "__main__":
    try:
        main()
    except Exception:
        print("❌ 程式發生未處理例外：")
        traceback.print_exc()
        # ✅ 不暫停，直接回 CMD
