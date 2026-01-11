# 檔名：P爬文章批次轉成docx.py
import os
import re
import csv
import time
import traceback
from io import BytesIO
from datetime import datetime
from urllib.parse import urljoin, urlparse

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from docx.image.exceptions import UnrecognizedImageError

try:
    from PIL import Image
    PIL_OK = True
except Exception:
    PIL_OK = False

CSV_PATH = r"F:\F\AI\web\web.csv"
OUT_DIR  = r"F:\F\AI\web"
DEBUG_DIR = os.path.join(OUT_DIR, "_debug")

SLEEP_SEC = 0.5


def safe_filename(name: str, max_len: int = 120) -> str:
    name = re.sub(r'[<>:"/\\|?*]', "_", (name or "")).strip()
    name = re.sub(r"\s+", " ", name)
    if len(name) > max_len:
        name = name[:max_len].rstrip()
    return name or ""


def read_urls_from_csv(csv_path: str):
    rows = []
    with open(csv_path, "r", encoding="utf-8-sig", newline="") as f:
        reader = csv.reader(f)
        for r in reader:
            if not r:
                continue
            url = (r[0] if len(r) > 0 else "").strip()
            name = (r[1] if len(r) > 1 else "").strip()

            if url.lower() in {"url", "網址"}:
                continue
            if not url:
                continue

            rows.append((url, name))
    return rows


def fetch_html(session: requests.Session, url: str) -> str:
    url = url.split("#", 1)[0]
    r = session.get(url, timeout=30, allow_redirects=True)
    r.raise_for_status()
    if not r.encoding or r.encoding.lower() == "iso-8859-1":
        r.encoding = r.apparent_encoding or "utf-8"
    return r.text


def looks_like_pixnet_post(html: str, url: str) -> bool:
    # Pixnet 文章頁通常會包含 blog/posts/<id>
    if "/blog/posts/" in url:
        # HTML 裡如果完全沒有 posts id 或 post 相關結構，可能被導到別頁
        if re.search(r"/blog/posts/\d+", html):
            return True
        # 退而求其次：出現 pixnet 文章常見字樣
        if "pixnet" in html.lower() and ("文章" in html or "發表" in html):
            return True
        return False
    return True


def extract_title(soup: BeautifulSoup) -> str:
    # ✅ 優先從文章內容抓 title（避免 <title> 被首頁/錯誤頁干擾）
    selectors = [
        "h1.title",
        "h1.post-title",
        "h1",
        "article h1",
        ".post-title",
        ".title",
        "h2.title",
        "h2",
    ]
    for sel in selectors:
        el = soup.select_one(sel)
        if el:
            t = el.get_text(" ", strip=True)
            if t and len(t) >= 2:
                return t

    if soup.title:
        t = soup.title.get_text(" ", strip=True)
        if t:
            # 常見形如「XXX @ 某某的部落格 :: 痞客邦」
            t = re.split(r"@|::", t)[0].strip()
            return t

    return "output"


def _parse_date_to_yyyymmdd(s: str):
    if not s:
        return None
    s = s.strip()

    m = re.search(r"(20\d{2})-(\d{1,2})-(\d{1,2})", s)
    if m:
        y, mo, d = m.group(1), int(m.group(2)), int(m.group(3))
        return f"{y}{mo:02d}{d:02d}"

    m = re.search(r"(20\d{2})/(\d{1,2})/(\d{1,2})", s)
    if m:
        y, mo, d = m.group(1), int(m.group(2)), int(m.group(3))
        return f"{y}{mo:02d}{d:02d}"

    return None


def extract_date8(soup: BeautifulSoup, html: str) -> str:
    # 1) meta
    meta_keys = [
        ("property", "article:published_time"),
        ("property", "og:published_time"),
        ("name", "pubdate"),
        ("name", "publishdate"),
        ("name", "publish_date"),
        ("name", "date"),
        ("itemprop", "datePublished"),
    ]
    for attr, val in meta_keys:
        tag = soup.find("meta", attrs={attr: val})
        if tag and tag.get("content"):
            d8 = _parse_date_to_yyyymmdd(tag["content"])
            if d8:
                return d8

    # 2) time
    t = soup.find("time")
    if t:
        dt = t.get("datetime") or t.get_text(" ", strip=True)
        d8 = _parse_date_to_yyyymmdd(dt)
        if d8:
            return d8

    # 3) Pixnet 常見：內文會出現 yyyy/mm/dd（甚至「發表於 2023/10/11」）
    m = re.search(r"(20\d{2})[/-](\d{1,2})[/-](\d{1,2})", html)
    if m:
        y, mo, d = m.group(1), int(m.group(2)), int(m.group(3))
        return f"{y}{mo:02d}{d:02d}"

    # 4) 真的抓不到就用今天（但我們會 debug）
    return datetime.now().strftime("%Y%m%d")


# ===== 以下正文抽取與圖片處理：維持你單次版的做法（略）=====
def _node_score(node) -> int:
    if not node:
        return -10**9
    txt = node.get_text(" ", strip=True)
    tlen = len(txt)
    p = len(node.find_all("p"))
    li = len(node.find_all("li"))
    h = len(node.find_all(["h1", "h2", "h3", "h4"]))
    bad = 0
    for bad_sel in ["nav", "header", "footer", "aside"]:
        bad += len(node.find_all(bad_sel))
    score = min(tlen, 20000) + p * 300 + li * 120 + h * 200 - bad * 500
    return score


def pick_content_root(soup: BeautifulSoup):
    selectors = [
        "article", "main",
        ".post-content", ".entry-content", ".article-content",
        ".content", "#content", "body"
    ]
    candidates = []
    for sel in selectors:
        candidates.extend(soup.select(sel))
    if not candidates:
        return soup.body or soup
    return max(candidates, key=_node_score)


def is_probably_nav_or_junk(tag) -> bool:
    if tag.name in {"nav", "header", "footer", "aside", "script", "style", "noscript"}:
        return True
    cls = " ".join(tag.get("class", [])).lower()
    if any(k in cls for k in ["share", "related", "sidebar", "widget", "comment", "ads", "advert", "breadcrumb"]):
        return True
    return False


def iter_content_blocks(root):
    for t in root.find_all(["script", "style", "noscript"]):
        t.decompose()

    for el in root.find_all(["h1", "h2", "h3", "h4", "p", "li", "blockquote", "pre", "img"]):
        if is_probably_nav_or_junk(el):
            continue

        if el.name in ["h1", "h2", "h3", "h4"]:
            txt = el.get_text(" ", strip=True)
            if txt:
                yield ("heading", el.name, txt)
            continue

        if el.name == "p":
            txt = el.get_text(" ", strip=True)
            if txt:
                yield ("p", txt)
            continue

        if el.name == "li":
            txt = el.get_text(" ", strip=True)
            if txt:
                yield ("li", txt)
            continue

        if el.name == "blockquote":
            txt = el.get_text(" ", strip=True)
            if txt:
                yield ("quote", txt)
            continue

        if el.name == "pre":
            txt = el.get_text("\n", strip=True)
            if txt:
                yield ("codeblock", txt)
            continue

        if el.name == "img":
            src = el.get("src") or el.get("data-src") or el.get("data-lazy-src") or el.get("data-original")
            if not src:
                continue
            alt = (el.get("alt") or "").strip()
            yield ("img", src, alt)
            continue


def download_image(session: requests.Session, img_url: str):
    try:
        r = session.get(img_url, timeout=30)
        r.raise_for_status()
        ctype = (r.headers.get("Content-Type") or "").lower()
        if "image" not in ctype:
            return None, ctype
        return r.content, ctype
    except Exception:
        return None, ""


def maybe_convert_webp_to_png_bytes(img_bytes: bytes, ctype: str, img_url: str):
    low_ct = (ctype or "").lower()
    ext = os.path.splitext(urlparse(img_url).path.lower())[1]
    is_webp = ("image/webp" in low_ct) or (ext == ".webp")
    if not is_webp or not PIL_OK:
        return None
    try:
        im = Image.open(BytesIO(img_bytes))
        out = BytesIO()
        im.convert("RGB").save(out, format="PNG")
        return out.getvalue()
    except Exception:
        return None


def write_docx_from_blocks(title: str, url: str, date8: str, blocks, session: requests.Session, base_url: str, out_path: str):
    doc = Document()
    doc.add_heading(title or "article", level=0)
    doc.add_paragraph(f"來源網址：{url}")
    doc.add_paragraph(f"建檔日期：{date8}")
    doc.add_paragraph("")

    for block in blocks:
        kind = block[0]
        if kind == "heading":
            _, tagname, txt = block
            level_map = {"h1": 1, "h2": 2, "h3": 3, "h4": 4}
            doc.add_heading(txt, level=level_map.get(tagname, 2))
        elif kind == "p":
            _, txt = block
            doc.add_paragraph(txt)
        elif kind == "li":
            _, txt = block
            doc.add_paragraph(txt, style="List Bullet")
        elif kind == "quote":
            _, txt = block
            doc.add_paragraph(txt, style="Intense Quote")
        elif kind == "codeblock":
            _, txt = block
            p = doc.add_paragraph()
            run = p.add_run(txt)
            run.font.name = "Consolas"
        elif kind == "img":
            _, src, alt = block
            img_url = urljoin(base_url.split("#", 1)[0], src)
            pth = urlparse(img_url).path.lower()
            if any(pth.endswith(x) for x in [".svg", ".ico"]):
                continue

            img, ctype = download_image(session, img_url)
            if not img:
                continue

            if alt:
                doc.add_paragraph(alt)

            try:
                converted = maybe_convert_webp_to_png_bytes(img, ctype, img_url)
                if converted:
                    doc.add_picture(BytesIO(converted), width=Inches(6.0))
                else:
                    doc.add_picture(BytesIO(img), width=Inches(6.0))
                time.sleep(SLEEP_SEC)
            except UnrecognizedImageError:
                print(f"⚠️ 無法識別圖片格式，跳過：{img_url}")
            except Exception as e:
                print(f"⚠️ 圖片插入失敗，跳過：{img_url} | {e}")

    doc.save(out_path)


def save_debug_html(url: str, html: str, idx: int):
    os.makedirs(DEBUG_DIR, exist_ok=True)
    fn = safe_filename(f"{idx}_{urlparse(url).path.split('/')[-1] or 'page'}.html") or f"{idx}.html"
    p = os.path.join(DEBUG_DIR, fn)
    with open(p, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"[DEBUG] 已存 HTML 供檢查：{p}")


def main():
    print("[START] 批次抓 Pixnet 文章 → docx（修正檔名/日期）")
    print(f"[INFO] CSV_PATH: {CSV_PATH}")
    print(f"[INFO] OUT_DIR : {OUT_DIR}")

    if not os.path.isfile(CSV_PATH):
        print(f"[ERROR] 找不到 CSV：{CSV_PATH}")
        return

    os.makedirs(OUT_DIR, exist_ok=True)

    items = read_urls_from_csv(CSV_PATH)
    print(f"[INFO] CSV 讀到 {len(items)} 筆網址")

    with requests.Session() as s:
        s.headers.update({
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8",
            "Accept-Language": "zh-TW,zh;q=0.9,en;q=0.6",
            "Referer": "https://www.google.com/"
        })

        ok = skip = fail = 0

        for idx, (url, name_in_csv) in enumerate(items, start=1):
            print(f"\n[DO] ({idx}/{len(items)}) {url}")

            try:
                html = fetch_html(s, url)

                # 如果抓到的不像文章頁，先存 debug 讓你看原因
                if not looks_like_pixnet_post(html, url):
                    save_debug_html(url, html, idx)

                soup = BeautifulSoup(html, "lxml")

                title = extract_title(soup)
                date8 = extract_date8(soup, html)

                # ✅ 檔名：CSV B欄優先；否則用文章標題；再不行才 output
                base = safe_filename(name_in_csv) or safe_filename(title) or "output"

                out_path = os.path.join(OUT_DIR, f"{date8}_{base}.docx")

                if os.path.exists(out_path):
                    print(f"[SKIP] 已存在：{os.path.basename(out_path)}")
                    skip += 1
                    continue

                root = pick_content_root(soup)
                blocks = list(iter_content_blocks(root))

                # 若正文太少，也存 debug（通常是被導去別頁或防爬）
                text_len = len(root.get_text(" ", strip=True))
                if text_len < 200:
                    save_debug_html(url, html, idx)

                write_docx_from_blocks(
                    title=title,
                    url=url,
                    date8=date8,
                    blocks=blocks,
                    session=s,
                    base_url=url,
                    out_path=out_path
                )

                print(f"[OK]  輸出：{os.path.basename(out_path)}")
                ok += 1

            except Exception as e:
                print(f"[ERR] 失敗：{e}")
                fail += 1

        print(f"\n[DONE] OK={ok}, SKIP={skip}, FAIL={fail}")


if __name__ == "__main__":
    try:
        main()
    except Exception:
        print("[FATAL] 未捕捉例外：")
        traceback.print_exc()