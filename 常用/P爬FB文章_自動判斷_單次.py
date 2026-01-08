# -*- coding: utf-8 -*-
"""
FB 貼文/相片頁擷取（最終版 - 全自動無互動）
- 支援 share/、permalink、photo.php 等
- 去除 FB UI 雜訊，只保留真正內文段落
- 圖片直接嵌入 DOCX（不存檔）
- DOCX 第一行顯示來源（社團/頁面）
- 頁面邊界：上下左右 1.5 cm
- 檔名：YYYYMMDD_文章第一行.docx（若被占用自動加 _01/_02）
- 依來源名稱建資料夾
- ✅ 無需任何 Enter 確認，爬完自動結束回到 CMD
"""

import re
import time
from io import BytesIO
from pathlib import Path
from urllib.parse import urlsplit, urlunsplit
from datetime import datetime

from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Pt, Inches, Cm
from PIL import Image


# ========= 基本設定 =========
BASE_DIR = Path(r"F:\F\AI")
PROFILE_DIR = BASE_DIR / "_fb_profile"


# ========= 工具 =========
def normalize_url(u: str) -> str:
    sp = urlsplit((u or "").strip())
    return urlunsplit((sp.scheme, sp.netloc, sp.path, sp.query, ""))


def safe_filename(s: str, max_len=60) -> str:
    s = re.sub(r'[<>:"/\\|?*]', "_", s or "").strip()
    s = re.sub(r"\s+", " ", s)
    s = s[:max_len].strip()
    return s or "FB內容"


def choose_available_path(folder: Path, base_name: str) -> Path:
    p0 = folder / f"{base_name}.docx"
    if not p0.exists():
        return p0
    for i in range(1, 200):
        p = folder / f"{base_name}_{i:02d}.docx"
        if not p.exists():
            return p
    return folder / f"{base_name}_{int(time.time())}.docx"


def scroll_to_load(page, times=8):
    for _ in range(times):
        try:
            page.mouse.wheel(0, 1800)
            page.wait_for_timeout(650)
        except Exception:
            pass


def wait_stable(page, total_ms=8000):
    """等 FB 跳轉/載入穩定（share/短連結很常跳）"""
    end = time.time() + total_ms / 1000
    last = ""
    while time.time() < end:
        try:
            cur = page.url
        except Exception:
            cur = ""
        if cur == last and cur:
            break
        last = cur
        try:
            page.wait_for_timeout(500)
        except Exception:
            pass


# ========= UI 雜訊過濾 =========
_UI_PATTERNS = [
    r"^讚$|^留言$|^分享$|^回覆$|^查看更多$|^更多$",
    r"^Like$|^Comment$|^Share$|^Reply$|^See more$",
    r"^最相關$|^最新$|^Top comments$|^Most relevant$",
    r"查看翻譯|See translation|翻譯",
    r"^.*人讚$|^.*個讚$|^.*則留言$|^.*次分享$",
]

def is_ui_noise(line: str) -> bool:
    t = (line or "").strip()
    if not t or len(t) < 3:
        return True
    for p in _UI_PATTERNS:
        if re.search(p, t, flags=re.IGNORECASE):
            return True
    return False


def clean_post_text(raw: str) -> str:
    if not raw:
        return ""
    raw = raw.replace("\r\n", "\n").replace("\r", "\n")
    lines = [l.strip() for l in raw.split("\n")]
    lines = [l for l in lines if l and not is_ui_noise(l)]

    seen = set()
    out = []
    for l in lines:
        k = re.sub(r"\s+", " ", l)
        if k not in seen:
            seen.add(k)
            out.append(l)

    if out and len(out[0]) <= 8:
        out = out[1:]

    return "\n".join(out).strip()


def first_content_line(text: str) -> str:
    for l in (text or "").splitlines():
        l = l.strip()
        if l:
            return l
    return ""


# ========= 來源/社團名稱 =========
def get_source_name(page) -> str:
    for sel in ['a[href*="/groups/"] span', 'a[href*="/groups/"]', 'h1', 'title']:
        try:
            loc = page.locator(sel)
            if loc.count():
                t = loc.first.inner_text(timeout=1500).strip()
                if t and len(t) <= 80:
                    return t
        except Exception:
            pass
    return "Facebook"


# ========= Dialog 判斷（排除 Messenger） =========
def find_best_dialog(page):
    dialogs = page.locator('div[role="dialog"]')
    best = None
    best_score = -1

    for i in range(dialogs.count()):
        d = dialogs.nth(i)
        aria = (d.get_attribute("aria-label") or "").lower()
        if "messenger" in aria or "chat" in aria:
            continue

        score = 0
        try:
            if d.locator('div[data-ad-preview="message"]').count() > 0:
                score += 6
        except Exception:
            pass
        try:
            if d.locator('img[data-visualcompletion="media-vc-image"]').count() > 0:
                score += 5
        except Exception:
            pass
        try:
            score += min(d.locator('div[dir="auto"]').count(), 10)
        except Exception:
            pass

        if score > best_score:
            best_score = score
            best = d

    return best if best_score >= 3 else None


def pick_text_container(container, page):
    selectors = [
        'div[data-ad-preview="message"]',
        'div[data-testid="post_message"]',
        'div[data-testid="photo-caption"]',
        'div[aria-label="相片說明"]',
        'div[aria-label="Photo caption"]',
    ]
    for sel in selectors:
        try:
            loc = container.locator(sel)
            if loc.count():
                return loc.first
        except Exception:
            pass

    try:
        main = page.locator('div[role="main"]').first
        for sel in selectors:
            loc = main.locator(sel)
            if loc.count():
                return loc.first
    except Exception:
        pass

    return container


# ========= 抓圖片（bytes，不落地） =========
def collect_image_bytes(ctx, container, page):
    url_list = []
    seen = set()

    def add_url(u):
        if not u or not u.startswith("http"):
            return
        lu = u.lower()
        if any(x in lu for x in ("emoji", "sprite", "static.xx.fbcdn.net")):
            return
        if u not in seen:
            seen.add(u)
            url_list.append(u)

    try:
        big = page.locator('img[data-visualcompletion="media-vc-image"]')
        for i in range(min(big.count(), 10)):
            add_url(big.nth(i).get_attribute("src") or "")
    except Exception:
        pass

    try:
        imgs = container.locator("img")
        for i in range(min(imgs.count(), 150)):
            add_url(imgs.nth(i).get_attribute("src") or "")
    except Exception:
        pass

    images = []
    for u in url_list:
        try:
            r = ctx.request.get(u, timeout=60000)
            if not r.ok:
                continue
            im = Image.open(BytesIO(r.body())).convert("RGB")
            bio = BytesIO()
            im.save(bio, format="JPEG", quality=92)
            bio.seek(0)
            images.append(bio)
        except Exception:
            pass

    return images


# ========= 抽內容 =========
def extract_fb_content(page, ctx):
    time.sleep(1)

    dlg = find_best_dialog(page)
    if dlg:
        container = dlg
        mode = "dialog"
    else:
        mode = "page"
        container = page.locator('div[role="article"]').first
        if container.count() == 0:
            container = page.locator('div[role="main"]').first

    # 標題
    title = ""
    try:
        title = (page.locator('meta[property="og:title"]').get_attribute("content") or "").strip()
    except Exception:
        pass
    if not title:
        try:
            title = page.title() or "Facebook 內容"
        except Exception:
            title = "Facebook 內容"

    # 作者
    author = ""
    try:
        a = container.locator("strong a")
        if a.count():
            author = a.first.inner_text(timeout=1500).strip()
    except Exception:
        pass

    text_container = pick_text_container(container, page)
    raw = ""
    try:
        raw = text_container.inner_text(timeout=4500)
    except Exception:
        try:
            raw = container.inner_text(timeout=4500)
        except Exception:
            raw = ""

    content = clean_post_text(raw)
    images = collect_image_bytes(ctx, container, page)

    return {
        "模式": mode,
        "標題": title,
        "作者": author,
        "內容": content,
        "圖片bytes": images,
        "網址": page.url,
    }


# ========= DOCX =========
def export_docx(data, source_name):
    today = datetime.now().strftime("%Y%m%d")
    folder = BASE_DIR / safe_filename(source_name, max_len=60)
    folder.mkdir(parents=True, exist_ok=True)

    base = first_content_line(data.get("內容", "")) or data.get("標題", "Facebook內容")
    base_name = f"{today}_{safe_filename(base, max_len=80)}"

    out_path = choose_available_path(folder, base_name)

    doc = Document()

    sec = doc.sections[0]
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(1.5)
    sec.right_margin = Cm(1.5)

    p = doc.add_paragraph(f"【來源】{source_name}")
    if p.runs:
        p.runs[0].bold = True

    h = doc.add_heading(data.get("標題", "Facebook 內容"), level=1)
    if h.runs:
        h.runs[0].font.size = Pt(16)

    if data.get("作者"):
        doc.add_paragraph(f"作者：{data['作者']}")
    doc.add_paragraph(f"建檔日期：{today}")
    doc.add_paragraph(f"原始網址：{data.get('網址','')}")
    doc.add_paragraph("")

    body = (data.get("內容") or "").strip()
    if body:
        for line in body.splitlines():
            doc.add_paragraph(line)
    else:
        doc.add_paragraph("（本頁面未抓到文字內文；可能只有圖片或權限限制）")

    for bio in data.get("圖片bytes", []):
        try:
            doc.add_paragraph("")
            doc.add_picture(bio, width=Inches(5.8))
        except Exception:
            pass

    # 若被鎖檔，自動換名重存
    try:
        doc.save(out_path)
        return out_path
    except PermissionError:
        for i in range(1, 11):
            alt = out_path.with_name(out_path.stem + f"_LOCK{i:02d}.docx")
            try:
                doc.save(alt)
                return alt
            except PermissionError:
                time.sleep(0.3)
        raise


# ========= 主程式（全自動） =========
def main():
    url = normalize_url(input("請輸入 FB 貼文網址：\n"))

    BASE_DIR.mkdir(parents=True, exist_ok=True)
    PROFILE_DIR.mkdir(parents=True, exist_ok=True)

    with sync_playwright() as p:
        ctx = p.chromium.launch_persistent_context(
            user_data_dir=str(PROFILE_DIR),
            headless=False,
            locale="zh-TW",
            viewport={"width": 1280, "height": 900},
        )
        page = ctx.new_page()

        page.goto(url, wait_until="domcontentloaded", timeout=60000)
        wait_stable(page, total_ms=10000)

        # 多捲幾次，讓 share/相片頁/貼文內容載入
        scroll_to_load(page, times=10)
        wait_stable(page, total_ms=6000)

        source = get_source_name(page)
        data = extract_fb_content(page, ctx)
        out = export_docx(data, source)

        print("✅ 已輸出：", out)
        print(f"模式：{data['模式']}｜圖片嵌入：{len(data.get('圖片bytes', []))} 張")

        # ✅ 自動關閉並回到 CMD
        ctx.close()


if __name__ == "__main__":
    main()
